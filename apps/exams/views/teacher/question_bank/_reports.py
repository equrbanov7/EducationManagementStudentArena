"""question_bank paketi — hesabat (xlsx/docx) qurucuları."""

from io import BytesIO

from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from django.utils.translation import pgettext

from ._helpers import (
    _format_int_list,
    _question_bank_feedback,
    _question_bank_source_diagnostics,
    _question_bank_warning_label,
    _warning_reference_text,
)

_EXCEL_SHEET_FORBIDDEN_CHARS = frozenset(":\\/?*[]")


def _tx(message):
    return pgettext("exams.export", message)


def _excel_sheet_title(message):
    title = _tx(message)
    title = "".join(" " if char in _EXCEL_SHEET_FORBIDDEN_CHARS else char for char in title)
    title = " ".join(title.split()).strip("'")
    return (title or message)[:31]


def _build_question_bank_report_xlsx(
    *,
    exam,
    raw_text,
    parsed,
    test_level_warnings,
    category_counts,
    warning_count,
    duplicate_count,
    error_count,
):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return _build_question_bank_report_docx(
            exam=exam,
            raw_text=raw_text,
            parsed=parsed,
            test_level_warnings=test_level_warnings,
            category_counts=category_counts,
            warning_count=warning_count,
            duplicate_count=duplicate_count,
            error_count=error_count,
        )

    parsed = parsed or []
    test_level_warnings = test_level_warnings or []
    diagnostics = _question_bank_source_diagnostics(raw_text, parsed)

    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = _excel_sheet_title("Xülasə")
    ws_problems = wb.create_sheet(_excel_sheet_title("Problemlər"))

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    section_font = Font(bold=True, color="1F2937")
    wrap = Alignment(wrap_text=True, vertical="top")

    summary_rows = [
        (_tx("İmtahan"), exam.title),
        (
            _tx("Hesabat vaxtı"),
            timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M"),
        ),
        (_tx("Parse olunan sual sayı"), len(parsed)),
        (_tx("Mənbədə END_QUESTION blokları"), diagnostics["source_block_count"]),
        (_tx("Dolu mənbə blokları"), diagnostics["nonempty_block_count"]),
        (
            _tx("Boş mənbə blokları"),
            _format_int_list(diagnostics["empty_blocks"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Parse olunmayan dolu blok sayı"),
            diagnostics["dropped_nonempty_block_count"],
        ),
        (
            _tx("Yazılı görünən ən böyük sual nömrəsi"),
            diagnostics["max_declared_number"] or "",
        ),
        (
            _tx("Ən böyük nömrə ilə parse sayı fərqi"),
            diagnostics["gap_from_max_declared"] or 0,
        ),
        (
            _tx("Nömrəsi görünməyən bloklar"),
            _format_int_list(diagnostics["no_visible_number_blocks"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Yazılı nömrələrdə boşluqlar"),
            _format_int_list(diagnostics["missing_numbers"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Təkrar yazılı nömrələr"),
            _format_int_list(diagnostics["duplicate_numbers"]) or _tx("Yoxdur"),
        ),
        (_tx("Xətalı sual sayı"), error_count),
        (_tx("Xəbərdarlıq sayı"), warning_count),
        (_tx("Dublikat sual sayı"), duplicate_count),
        (_tx("Variant problemi"), (category_counts or {}).get("structure", 0)),
        (_tx("Uzunluq balansı"), (category_counts or {}).get("balance", 0)),
        (_tx("Təmiz sual sayı"), (category_counts or {}).get("clean", 0)),
    ]

    ws_summary.append([_tx("Sahə"), _tx("Dəyər")])
    for cell in ws_summary[1]:
        cell.fill = header_fill
        cell.font = header_font
    for row in summary_rows:
        ws_summary.append(row)

    ws_summary.append([])
    ws_summary.append(
        [
            _tx("Qeyd"),
            _tx("Feedback sütunu müəllimə göndərmək üçün hazır düzəliş mətnidir."),
        ]
    )
    ws_summary.cell(row=ws_summary.max_row, column=1).font = section_font

    for row in ws_summary.iter_rows():
        for cell in row:
            cell.alignment = wrap
    ws_summary.column_dimensions["A"].width = 34
    ws_summary.column_dimensions["B"].width = 86

    problem_headers = [
        _tx("Sual sıra #"),
        _tx("Mənbə nömrəsi"),
        _tx("Səviyyə"),
        _tx("Problem növü"),
        _tx("Problem izahı"),
        _tx("Bağlı sual/ref"),
        _tx("Feedback"),
        _tx("Sual mətni"),
        "A",
        "B",
        "C",
        "D",
        "E",
        _tx("Düz cavab"),
    ]
    ws_problems.append(problem_headers)
    for cell in ws_problems[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = wrap

    row_count = 0
    for warning in test_level_warnings:
        warning_type = warning.get("type")
        ws_problems.append(
            [
                _tx("Test ümumi"),
                "",
                warning.get("severity", "warning"),
                _question_bank_warning_label(warning_type),
                warning.get("msg", ""),
                "",
                _question_bank_feedback(warning_type),
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ]
        )
        row_count += 1

    for idx, q in enumerate(parsed, start=1):
        warnings = q.get("warnings") or []
        if not warnings:
            continue
        opts = q.get("options") or {}
        for warning in warnings:
            warning_type = warning.get("type")
            ws_problems.append(
                [
                    idx,
                    q.get("q_no", ""),
                    warning.get("severity", "warning"),
                    _question_bank_warning_label(warning_type),
                    warning.get("msg", ""),
                    _warning_reference_text(warning),
                    _question_bank_feedback(warning_type),
                    q.get("text", ""),
                    opts.get("A", ""),
                    opts.get("B", ""),
                    opts.get("C", ""),
                    opts.get("D", ""),
                    opts.get("E", ""),
                    ", ".join(q.get("correct") or []),
                ]
            )
            row_count += 1

    if row_count == 0:
        ws_problems.append(
            [
                "",
                "",
                "info",
                _tx("Problem yoxdur"),
                _tx("Import edilmiş suallarda xəbərdarlıq tapılmadı."),
                "",
                _tx("Düzəliş tələb olunmur."),
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ]
        )

    widths = [12, 14, 12, 22, 60, 24, 58, 70, 34, 34, 34, 34, 34, 14]
    for idx, width in enumerate(widths, start=1):
        ws_problems.column_dimensions[get_column_letter(idx)].width = width
    for row in ws_problems.iter_rows():
        for cell in row:
            cell.alignment = wrap

    ws_problems.freeze_panes = "A2"
    ws_summary.freeze_panes = "A2"

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    filename = (
        f"{slugify(exam.title) or 'test-bank'}-problem-report-"
        f"{timezone.localtime(timezone.now()).strftime('%Y%m%d-%H%M')}.xlsx"
    )
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _build_question_bank_report_docx(
    *,
    exam,
    raw_text,
    parsed,
    test_level_warnings,
    category_counts,
    warning_count,
    duplicate_count,
    error_count,
):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Add python-docx to requirements.") from exc

    parsed = parsed or []
    test_level_warnings = test_level_warnings or []
    diagnostics = _question_bank_source_diagnostics(raw_text, parsed)

    doc = Document()
    doc.add_heading(
        _tx("%(title)s - problem hesabatı") % {"title": exam.title},
        level=1,
    )
    doc.add_paragraph(
        _tx("Hesabat vaxtı: %(datetime)s") % {"datetime": timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")}
    )

    doc.add_heading(_tx("Xülasə"), level=2)
    summary_rows = [
        (_tx("Parse olunan sual sayı"), len(parsed)),
        (_tx("Mənbədə END_QUESTION blokları"), diagnostics["source_block_count"]),
        (_tx("Dolu mənbə blokları"), diagnostics["nonempty_block_count"]),
        (
            _tx("Boş mənbə blokları"),
            _format_int_list(diagnostics["empty_blocks"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Parse olunmayan dolu blok sayı"),
            diagnostics["dropped_nonempty_block_count"],
        ),
        (
            _tx("Yazılı görünən ən böyük sual nömrəsi"),
            diagnostics["max_declared_number"] or "",
        ),
        (
            _tx("Ən böyük nömrə ilə parse sayı fərqi"),
            diagnostics["gap_from_max_declared"] or 0,
        ),
        (
            _tx("Nömrəsi görünməyən bloklar"),
            _format_int_list(diagnostics["no_visible_number_blocks"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Yazılı nömrələrdə boşluqlar"),
            _format_int_list(diagnostics["missing_numbers"]) or _tx("Yoxdur"),
        ),
        (
            _tx("Təkrar yazılı nömrələr"),
            _format_int_list(diagnostics["duplicate_numbers"]) or _tx("Yoxdur"),
        ),
        (_tx("Xətalı sual sayı"), error_count),
        (_tx("Xəbərdarlıq sayı"), warning_count),
        (_tx("Dublikat sual sayı"), duplicate_count),
        (_tx("Variant problemi"), (category_counts or {}).get("structure", 0)),
        (_tx("Uzunluq balansı"), (category_counts or {}).get("balance", 0)),
        (_tx("Təmiz sual sayı"), (category_counts or {}).get("clean", 0)),
    ]
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = _tx("Sahə")
    summary_table.rows[0].cells[1].text = _tx("Dəyər")
    for label, value in summary_rows:
        cells = summary_table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    doc.add_paragraph(_tx("Feedback qeydləri müəllimə göndərmək üçün hazır düzəliş mətnidir."))
    doc.add_heading(_tx("Problemlər"), level=2)

    problem_count = 0
    for warning in test_level_warnings:
        warning_type = warning.get("type")
        paragraph = doc.add_paragraph()
        paragraph.add_run(_tx("Test ümumi")).bold = True
        paragraph.add_run("\n" + _tx("Səviyyə: %(severity)s") % {"severity": warning.get("severity", "warning")})
        paragraph.add_run("\n" + _tx("Problem növü: %(type)s") % {"type": _question_bank_warning_label(warning_type)})
        paragraph.add_run("\n" + _tx("Problem: %(problem)s") % {"problem": warning.get("msg", "")})
        paragraph.add_run("\n" + _tx("Feedback: %(feedback)s") % {"feedback": _question_bank_feedback(warning_type)})
        problem_count += 1

    for idx, q in enumerate(parsed, start=1):
        opts = q.get("options") or {}
        for warning in q.get("warnings") or []:
            warning_type = warning.get("type")
            paragraph = doc.add_paragraph()
            paragraph.add_run(_tx("Sual #%(number)s") % {"number": idx}).bold = True
            if q.get("q_no"):
                paragraph.add_run(" | " + _tx("Mənbə nömrəsi: %(number)s") % {"number": q.get("q_no")})
            paragraph.add_run("\n" + _tx("Səviyyə: %(severity)s") % {"severity": warning.get("severity", "warning")})
            paragraph.add_run(
                "\n" + _tx("Problem növü: %(type)s") % {"type": _question_bank_warning_label(warning_type)}
            )
            paragraph.add_run("\n" + _tx("Problem: %(problem)s") % {"problem": warning.get("msg", "")})
            reference = _warning_reference_text(warning)
            if reference:
                paragraph.add_run("\n" + _tx("Bağlı sual/ref: %(reference)s") % {"reference": reference})
            paragraph.add_run(
                "\n" + _tx("Feedback: %(feedback)s") % {"feedback": _question_bank_feedback(warning_type)}
            )
            paragraph.add_run("\n" + _tx("Sual mətni: %(text)s") % {"text": q.get("text", "")})
            paragraph.add_run(
                "\n"
                + _tx("Variantlar: %(options)s")
                % {"options": "; ".join(f"{label}) {opts.get(label, '')}" for label in "ABCDE" if opts.get(label))}
            )
            paragraph.add_run("\n" + _tx("Düz cavab: %(answer)s") % {"answer": ", ".join(q.get("correct") or [])})
            problem_count += 1

    if problem_count == 0:
        paragraph = doc.add_paragraph()
        paragraph.add_run(_tx("Problem yoxdur")).bold = True
        paragraph.add_run("\n" + _tx("Import edilmiş suallarda xəbərdarlıq tapılmadı."))
        paragraph.add_run("\n" + _tx("Feedback: %(feedback)s") % {"feedback": _tx("Düzəliş tələb olunmur.")})

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = (
        f"{slugify(exam.title) or 'test-bank'}-problem-report-"
        f"{timezone.localtime(timezone.now()).strftime('%Y%m%d-%H%M')}.docx"
    )
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
