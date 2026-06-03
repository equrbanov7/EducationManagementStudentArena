import json
import logging
import re
from io import BytesIO
from urllib.parse import urlencode, urlsplit

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.db.models import Max
from django.http import Http404, HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.urls import Resolver404, resolve, reverse
from django.utils import timezone
from django.utils.http import url_has_allowed_host_and_scheme
from django.utils.text import slugify
from django.utils.translation import pgettext, pgettext_lazy
from django.views.decorators.http import require_POST

from apps.exams.constants import QUESTION_RE
from apps.exams.models import ExamQuestion, ExamQuestionOption, QuestionBlock
from apps.exams.services.access_policy import _ensure_teacher
from apps.exams.services.ai_question_generation import generate_question_bank_text
from apps.exams.services.coding_definition import sync_coding_questions_for_exam
from apps.exams.services.parsing import END_QUESTION_RE, extract_text_from_upload, parse_bulk_mcq
from apps.exams.services.utils import _norm
from apps.exams.views.shared.tenant import get_teacher_exam_or_404

WRITTEN_QUESTION_PREFIX_RE = re.compile(r"^\s*\d+\s*[\.\)]\s*", re.MULTILINE)
logger = logging.getLogger(__name__)


def _safe_same_origin_redirect_path(request, candidate_url):
    raw_url = (candidate_url or "").strip()
    if not raw_url:
        return ""

    if not url_has_allowed_host_and_scheme(
        raw_url,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        return ""

    parsed = urlsplit(raw_url)
    if parsed.netloc and parsed.netloc != request.get_host():
        return ""

    path = parsed.path or "/"
    query = f"?{parsed.query}" if parsed.query else ""
    fragment = f"#{parsed.fragment}" if parsed.fragment else ""
    return f"{path}{query}{fragment}"


def _is_internal_exam_management_path(candidate_path):
    raw_path = (candidate_path or "").strip()
    if not raw_path:
        return False

    try:
        match = resolve(urlsplit(raw_path).path)
    except Resolver404:
        return False

    return match.namespace == "exams" and match.url_name in {
        "teacher_questions_bank",
        "test_question_bank",
        "create_question_bank",
        "add_exam_question",
        "edit_exam_question",
        "delete_exam_question",
        "teacher_exam_results",
        "teacher_view_attempt",
        "teacher_check_attempt",
    }


def _resolve_question_bank_navigation(request):
    requested_profile_section = (request.GET.get("from_section") or request.POST.get("from_section") or "").strip()
    valid_profile_sections = {
        "my-exams",
        "assigned-exams",
        "profile-info",
        "my-courses",
        "assigned-courses",
        "courses",
        "pending-review",
        "review-results",
    }
    if requested_profile_section not in valid_profile_sections:
        requested_profile_section = ""

    return_to = _safe_same_origin_redirect_path(
        request,
        request.GET.get("return_to") or request.POST.get("return_to"),
    )
    if _is_internal_exam_management_path(return_to):
        return_to = ""

    nav_params = {}
    if requested_profile_section:
        nav_params["from_section"] = requested_profile_section
    if return_to:
        nav_params["return_to"] = return_to

    return requested_profile_section, return_to, urlencode(nav_params)


def _append_navigation_query(path, navigation_query):
    if not navigation_query:
        return path
    separator = "&" if "?" in path else "?"
    return f"{path}{separator}{navigation_query}"


def _parse_written_questions(content_text):
    text = (content_text or "").strip()
    if not text:
        return []

    matches = list(WRITTEN_QUESTION_PREFIX_RE.finditer(text))
    if not matches:
        return [text]

    questions = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        question_text = text[start:end].strip()
        if question_text:
            questions.append(question_text)

    return questions


def _question_bank_title_context(exam):
    if exam.exam_type == "coding":
        return {
            "question_bank_page_title": pgettext("exams.template.create_question_bank", "Praktiki sual bankı"),
            "question_bank_heading": pgettext("exams.template.create_question_bank", "praktiki sual bankı"),
        }

    return {
        "question_bank_page_title": pgettext("exams.template.create_question_bank", "Yazılı sual bankı"),
        "question_bank_heading": pgettext("exams.template.create_question_bank", "yazılı sual bankı"),
    }


def _optional_non_negative_int(value):
    raw = (value or "").strip()
    if not raw:
        return None
    try:
        parsed = int(raw)
    except (TypeError, ValueError):
        return None
    return parsed if parsed >= 0 else None


def _parse_selected_question_indices(post_data):
    if "selected_indices" in post_data:
        compact_value = (post_data.get("selected_indices") or "").strip()
        if compact_value:
            raw_values = compact_value.split(",")
        else:
            legacy_values = post_data.getlist("selected")
            if legacy_values:
                raw_values = legacy_values
            else:
                return set()
    else:
        raw_values = post_data.getlist("selected")
        if not raw_values:
            return None

    selected = set()
    for raw_value in raw_values:
        try:
            value = int(str(raw_value).strip())
        except (TypeError, ValueError):
            continue
        if value > 0:
            selected.add(value)
    return selected


def _parse_points_payload(post_data):
    raw_payload = (post_data.get("points_payload") or "").strip()
    if not raw_payload:
        return {}

    try:
        payload = json.loads(raw_payload)
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}

    if not isinstance(payload, dict):
        return {}
    return payload


def _sync_written_block_questions(block, question_texts):
    existing_questions = list(block.questions.order_by("order", "id"))
    normalized_texts = [text for text in question_texts if text]

    for index, q_text in enumerate(normalized_texts, start=1):
        if index <= len(existing_questions):
            question = existing_questions[index - 1]
            update_fields = []

            if question.text != q_text:
                question.text = q_text
                update_fields.append("text")
            if question.order != index:
                question.order = index
                update_fields.append("order")
            if question.block_id != block.id:
                question.block = block
                update_fields.append("block")
            if question.exam_id != block.exam_id:
                question.exam = block.exam
                update_fields.append("exam")
            if question.answer_mode != "single":
                question.answer_mode = "single"
                update_fields.append("answer_mode")

            if update_fields:
                question.save(update_fields=update_fields)
        else:
            ExamQuestion.objects.create(
                exam=block.exam,
                block=block,
                text=q_text,
                order=index,
                answer_mode="single",
            )

    for stale_question in existing_questions[len(normalized_texts) :]:
        stale_question.delete()


_WARNING_TYPE_LABELS = {
    "duplicate_in_import": "Dublikat",
    "already_in_exam": "Bankda dublikat",
    "duplicate_option_text": "Təkrar variant",
    "missing_option": "Variant çatışmır",
    "option_count_recommend_5": "Variant sayı",
    "option_count_too_low": "Variant sayı",
    "empty_option_text": "Boş variant",
    "correct_missing": "Düz cavab uyğunsuzluğu",
    "correct_too_long": "Uzunluq balansı",
    "correct_too_short": "Uzunluq balansı",
    "bulk_correct_too_long": "Test miqyaslı balans",
    "bulk_correct_too_short": "Test miqyaslı balans",
}


_WARNING_FEEDBACK = {
    "duplicate_in_import": (
        "Bu sual import faylında başqa sualla eynidir. Təkrarlanan suallardan birini silin "
        "və ya sual/variant mətnlərini fərqləndirin."
    ),
    "already_in_exam": (
        "Bu sual imtahan bankında artıq mövcuddur. Yeni sual kimi əlavə ediləcəksə mətni "
        "fərqləndirin, əks halda importdan çıxarın."
    ),
    "duplicate_option_text": "Eyni mətnli variantları dəyişin; hər variant ayrı məna daşımalıdır.",
    "missing_option": "A-D minimum variantlarını tamamlayın.",
    "option_count_recommend_5": "Mümkünsə E variantını əlavə edin ki, sual 5 variantlı standartla uyğun olsun.",
    "option_count_too_low": "Variant sayını ən azı 4-ə çatdırın.",
    "empty_option_text": "Boş variant mətnini doldurun və ya həmin variantı silib strukturu düzəldin.",
    "correct_missing": "Cavab sətrində göstərilən düzgün variant sualın variantları arasında olmalıdır.",
    "correct_too_long": (
        "Düzgün cavab digər variantlardan çox uzun görünür. Variantların uzunluğunu və "
        "detal səviyyəsini balanslaşdırın."
    ),
    "correct_too_short": (
        "Düzgün cavab digər variantlardan çox qısa görünür. Variantların uzunluğunu və "
        "detal səviyyəsini balanslaşdırın."
    ),
    "bulk_correct_too_long": "Test üzrə düzgün cavabların uzunluq patternini balanslaşdırın.",
    "bulk_correct_too_short": "Test üzrə düzgün cavabların qısalıq patternini balanslaşdırın.",
}


def _question_bank_warning_label(warning_type):
    return _WARNING_TYPE_LABELS.get(warning_type or "", warning_type or "Problem")


def _question_bank_feedback(warning_type):
    return _WARNING_FEEDBACK.get(
        warning_type or "",
        "Sual mətnini, variantları və düzgün cavab qeydini müəllim tərəfindən yenidən yoxlayın.",
    )


def _split_end_question_source_blocks(raw_text):
    blocks = []
    block = []

    for raw in (raw_text or "").splitlines():
        line = raw.strip()
        if END_QUESTION_RE.match(line):
            blocks.append(block)
            block = []
            continue
        if line:
            block.append(line)

    if block:
        blocks.append(block)

    return blocks


def _question_bank_source_diagnostics(raw_text, parsed):
    blocks = _split_end_question_source_blocks(raw_text)
    empty_blocks = [idx for idx, block in enumerate(blocks, start=1) if not block]
    nonempty_blocks = [block for block in blocks if block]

    declared = []
    no_visible_number_blocks = []
    for idx, block in enumerate(blocks, start=1):
        if not block:
            continue
        first_number = None
        for line in block:
            m_q = QUESTION_RE.match(line)
            if m_q:
                first_number = int(m_q.group(1))
                break
        if first_number is None:
            no_visible_number_blocks.append(idx)
        else:
            declared.append((idx, first_number))

    declared_numbers = [num for _, num in declared]
    duplicate_numbers = sorted({num for num in declared_numbers if declared_numbers.count(num) > 1})
    max_declared = max(declared_numbers) if declared_numbers else None
    missing_numbers = []
    if max_declared:
        visible_set = set(declared_numbers)
        missing_numbers = [num for num in range(1, max_declared + 1) if num not in visible_set]

    return {
        "source_block_count": len(blocks),
        "nonempty_block_count": len(nonempty_blocks),
        "empty_blocks": empty_blocks,
        "parsed_count": len(parsed or []),
        "dropped_nonempty_block_count": max(0, len(nonempty_blocks) - len(parsed or [])),
        "visible_number_count": len(declared_numbers),
        "max_declared_number": max_declared,
        "gap_from_max_declared": (max_declared - len(parsed or [])) if max_declared else None,
        "missing_numbers": missing_numbers,
        "duplicate_numbers": duplicate_numbers,
        "no_visible_number_blocks": no_visible_number_blocks,
    }


def _format_int_list(values, limit=40):
    values = list(values or [])
    if not values:
        return ""
    shown = ", ".join(str(v) for v in values[:limit])
    if len(values) > limit:
        shown += f" ... (+{len(values) - limit})"
    return shown


def _warning_reference_text(warning):
    warning_type = warning.get("type")
    if warning_type == "duplicate_in_import":
        refs = warning.get("all_refs") or []
        if not refs and warning.get("ref"):
            refs = [warning.get("ref")]
        return ", ".join(f"#{ref}" for ref in refs)
    if warning_type == "already_in_exam":
        parts = []
        if warning.get("ref_db_order"):
            parts.append(f"Bank sıra: {warning.get('ref_db_order')}")
        if warning.get("ref_db_id"):
            parts.append(f"ID: {warning.get('ref_db_id')}")
        return "; ".join(parts)
    return ""


def _build_question_bank_report_xlsx(
    *,
    exam,
    raw_text,
    parsed,
    test_level_warnings,
    category_counts,
    warning_count,
    duplicate_count,
    error_count,
):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return _build_question_bank_report_docx(
            exam=exam,
            raw_text=raw_text,
            parsed=parsed,
            test_level_warnings=test_level_warnings,
            category_counts=category_counts,
            warning_count=warning_count,
            duplicate_count=duplicate_count,
            error_count=error_count,
        )

    parsed = parsed or []
    test_level_warnings = test_level_warnings or []
    diagnostics = _question_bank_source_diagnostics(raw_text, parsed)

    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Xülasə"
    ws_problems = wb.create_sheet("Problemlər")

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    section_font = Font(bold=True, color="1F2937")
    wrap = Alignment(wrap_text=True, vertical="top")

    summary_rows = [
        ("İmtahan", exam.title),
        ("Hesabat vaxtı", timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")),
        ("Parse olunan sual sayı", len(parsed)),
        ("Mənbədə END_QUESTION blokları", diagnostics["source_block_count"]),
        ("Dolu mənbə blokları", diagnostics["nonempty_block_count"]),
        ("Boş mənbə blokları", _format_int_list(diagnostics["empty_blocks"]) or "Yoxdur"),
        ("Parse olunmayan dolu blok sayı", diagnostics["dropped_nonempty_block_count"]),
        ("Yazılı görünən ən böyük sual nömrəsi", diagnostics["max_declared_number"] or ""),
        ("Ən böyük nömrə ilə parse sayı fərqi", diagnostics["gap_from_max_declared"] or 0),
        ("Nömrəsi görünməyən bloklar", _format_int_list(diagnostics["no_visible_number_blocks"]) or "Yoxdur"),
        ("Yazılı nömrələrdə boşluqlar", _format_int_list(diagnostics["missing_numbers"]) or "Yoxdur"),
        ("Təkrar yazılı nömrələr", _format_int_list(diagnostics["duplicate_numbers"]) or "Yoxdur"),
        ("Xətalı sual sayı", error_count),
        ("Xəbərdarlıq sayı", warning_count),
        ("Dublikat sual sayı", duplicate_count),
        ("Variant problemi", (category_counts or {}).get("structure", 0)),
        ("Uzunluq balansı", (category_counts or {}).get("balance", 0)),
        ("Təmiz sual sayı", (category_counts or {}).get("clean", 0)),
    ]

    ws_summary.append(["Sahə", "Dəyər"])
    for cell in ws_summary[1]:
        cell.fill = header_fill
        cell.font = header_font
    for row in summary_rows:
        ws_summary.append(row)

    ws_summary.append([])
    ws_summary.append(["Qeyd", "Feedback sütunu müəllimə göndərmək üçün hazır düzəliş mətnidir."])
    ws_summary.cell(row=ws_summary.max_row, column=1).font = section_font

    for row in ws_summary.iter_rows():
        for cell in row:
            cell.alignment = wrap
    ws_summary.column_dimensions["A"].width = 34
    ws_summary.column_dimensions["B"].width = 86

    problem_headers = [
        "Sual sıra #",
        "Mənbə nömrəsi",
        "Səviyyə",
        "Problem növü",
        "Problem izahı",
        "Bağlı sual/ref",
        "Feedback",
        "Sual mətni",
        "A",
        "B",
        "C",
        "D",
        "E",
        "Düz cavab",
    ]
    ws_problems.append(problem_headers)
    for cell in ws_problems[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = wrap

    row_count = 0
    for warning in test_level_warnings:
        warning_type = warning.get("type")
        ws_problems.append(
            [
                "Test ümumi",
                "",
                warning.get("severity", "warning"),
                _question_bank_warning_label(warning_type),
                warning.get("msg", ""),
                "",
                _question_bank_feedback(warning_type),
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ]
        )
        row_count += 1

    for idx, q in enumerate(parsed, start=1):
        warnings = q.get("warnings") or []
        if not warnings:
            continue
        opts = q.get("options") or {}
        for warning in warnings:
            warning_type = warning.get("type")
            ws_problems.append(
                [
                    idx,
                    q.get("q_no", ""),
                    warning.get("severity", "warning"),
                    _question_bank_warning_label(warning_type),
                    warning.get("msg", ""),
                    _warning_reference_text(warning),
                    _question_bank_feedback(warning_type),
                    q.get("text", ""),
                    opts.get("A", ""),
                    opts.get("B", ""),
                    opts.get("C", ""),
                    opts.get("D", ""),
                    opts.get("E", ""),
                    ", ".join(q.get("correct") or []),
                ]
            )
            row_count += 1

    if row_count == 0:
        ws_problems.append(
            [
                "",
                "",
                "info",
                "Problem yoxdur",
                "Import edilmiş suallarda xəbərdarlıq tapılmadı.",
                "",
                "Düzəliş tələb olunmur.",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
            ]
        )

    widths = [12, 14, 12, 22, 60, 24, 58, 70, 34, 34, 34, 34, 34, 14]
    for idx, width in enumerate(widths, start=1):
        ws_problems.column_dimensions[get_column_letter(idx)].width = width
    for row in ws_problems.iter_rows():
        for cell in row:
            cell.alignment = wrap

    ws_problems.freeze_panes = "A2"
    ws_summary.freeze_panes = "A2"

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    filename = (
        f"{slugify(exam.title) or 'test-bank'}-problem-report-"
        f"{timezone.localtime(timezone.now()).strftime('%Y%m%d-%H%M')}.xlsx"
    )
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _build_question_bank_report_docx(
    *,
    exam,
    raw_text,
    parsed,
    test_level_warnings,
    category_counts,
    warning_count,
    duplicate_count,
    error_count,
):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Add python-docx to requirements.") from exc

    parsed = parsed or []
    test_level_warnings = test_level_warnings or []
    diagnostics = _question_bank_source_diagnostics(raw_text, parsed)

    doc = Document()
    doc.add_heading(f"{exam.title} - problem hesabatı", level=1)
    doc.add_paragraph(timezone.localtime(timezone.now()).strftime("Hesabat vaxtı: %Y-%m-%d %H:%M"))

    doc.add_heading("Xülasə", level=2)
    summary_rows = [
        ("Parse olunan sual sayı", len(parsed)),
        ("Mənbədə END_QUESTION blokları", diagnostics["source_block_count"]),
        ("Dolu mənbə blokları", diagnostics["nonempty_block_count"]),
        ("Boş mənbə blokları", _format_int_list(diagnostics["empty_blocks"]) or "Yoxdur"),
        ("Parse olunmayan dolu blok sayı", diagnostics["dropped_nonempty_block_count"]),
        ("Yazılı görünən ən böyük sual nömrəsi", diagnostics["max_declared_number"] or ""),
        ("Ən böyük nömrə ilə parse sayı fərqi", diagnostics["gap_from_max_declared"] or 0),
        ("Nömrəsi görünməyən bloklar", _format_int_list(diagnostics["no_visible_number_blocks"]) or "Yoxdur"),
        ("Yazılı nömrələrdə boşluqlar", _format_int_list(diagnostics["missing_numbers"]) or "Yoxdur"),
        ("Təkrar yazılı nömrələr", _format_int_list(diagnostics["duplicate_numbers"]) or "Yoxdur"),
        ("Xətalı sual sayı", error_count),
        ("Xəbərdarlıq sayı", warning_count),
        ("Dublikat sual sayı", duplicate_count),
        ("Variant problemi", (category_counts or {}).get("structure", 0)),
        ("Uzunluq balansı", (category_counts or {}).get("balance", 0)),
        ("Təmiz sual sayı", (category_counts or {}).get("clean", 0)),
    ]
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Sahə"
    summary_table.rows[0].cells[1].text = "Dəyər"
    for label, value in summary_rows:
        cells = summary_table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    doc.add_paragraph("Feedback qeydləri müəllimə göndərmək üçün hazır düzəliş mətnidir.")
    doc.add_heading("Problemlər", level=2)

    problem_count = 0
    for warning in test_level_warnings:
        warning_type = warning.get("type")
        paragraph = doc.add_paragraph()
        paragraph.add_run("Test ümumi").bold = True
        paragraph.add_run(f"\nSəviyyə: {warning.get('severity', 'warning')}")
        paragraph.add_run(f"\nProblem növü: {_question_bank_warning_label(warning_type)}")
        paragraph.add_run(f"\nProblem: {warning.get('msg', '')}")
        paragraph.add_run(f"\nFeedback: {_question_bank_feedback(warning_type)}")
        problem_count += 1

    for idx, q in enumerate(parsed, start=1):
        opts = q.get("options") or {}
        for warning in q.get("warnings") or []:
            warning_type = warning.get("type")
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"Sual #{idx}").bold = True
            if q.get("q_no"):
                paragraph.add_run(f" | Mənbə nömrəsi: {q.get('q_no')}")
            paragraph.add_run(f"\nSəviyyə: {warning.get('severity', 'warning')}")
            paragraph.add_run(f"\nProblem növü: {_question_bank_warning_label(warning_type)}")
            paragraph.add_run(f"\nProblem: {warning.get('msg', '')}")
            reference = _warning_reference_text(warning)
            if reference:
                paragraph.add_run(f"\nBağlı sual/ref: {reference}")
            paragraph.add_run(f"\nFeedback: {_question_bank_feedback(warning_type)}")
            paragraph.add_run(f"\nSual mətni: {q.get('text', '')}")
            paragraph.add_run(
                "\nVariantlar: " + "; ".join(f"{label}) {opts.get(label, '')}" for label in "ABCDE" if opts.get(label))
            )
            paragraph.add_run(f"\nDüz cavab: {', '.join(q.get('correct') or [])}")
            problem_count += 1

    if problem_count == 0:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Problem yoxdur").bold = True
        paragraph.add_run("\nImport edilmiş suallarda xəbərdarlıq tapılmadı.")
        paragraph.add_run("\nFeedback: Düzəliş tələb olunmur.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = (
        f"{slugify(exam.title) or 'test-bank'}-problem-report-"
        f"{timezone.localtime(timezone.now()).strftime('%Y%m%d-%H%M')}.docx"
    )
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


_QUESTION_BANK_TEMPLATE_TXT = """\
# EMSArena — Test sual bankı şablonu
# Hər sualın 4 və ya 5 variantı olmalıdır (A–E). Düz cavabı 3 üsuldan biri ilə qeyd edə bilərsiniz.
# 1) Cavabı sual sonunda yazın: "Cavab: B"
# 2) Düz variantın əvvəlinə * qoyun: "*B) Mətnim"
# 3) Heç bir işarə yoxdursa A standart olaraq düzgün sayılır.
# Çox cavablı suallarda virgül istifadə edin: "Cavab: A,C"
# Hər sualdan sonra boş sətir buraxın.
# Bu sətirlər (# ilə başlayanlar) və "END_QUESTION" işarəsi importda nəzərə alınmır.

1. Şəbəkədə məlumat hansı ölçü vahidi ilə ötürülür?
A) Bit
B) Bayt
C) Volt
D) Hertz
Cavab: A

2. Aşağıdakılardan hansı proqramlaşdırma dilidir?
A) Python
*B) JavaScript
*C) Java
D) Word
Cavab: A,B,C

3. HTML nədir?
A) Proqramlaşdırma dili
B) İşarələmə dili
C) Verilənlər bazası
D) Əməliyyat sistemi
*B)

4. OSI modelində ən aşağı səviyyə hansıdır?
A) Tətbiq səviyyəsi
B) Nəqliyyat səviyyəsi
*C) Fiziki səviyyə
D) Şəbəkə səviyyəsi
E) Heç biri
"""


@login_required
def test_question_bank_template_download(request, slug):
    """
    Müəllimə hazır sual bankı şablonu endirir.
    Format: ?format=txt (default) və ya ?format=docx
    """
    _ensure_teacher(request.user)
    # Slug yoxlaması — tenant izolyasiyası
    get_teacher_exam_or_404(request, slug=slug)

    file_format = (request.GET.get("format") or "txt").lower().strip()

    if file_format == "docx":
        try:
            from io import BytesIO

            from django.http import HttpResponse

            from docx import Document

            doc = Document()
            doc.add_heading("EMSArena — Test sual bankı şablonu", level=1)
            intro = doc.add_paragraph()
            intro.add_run(
                "Hər sualın 4 və ya 5 variantı olmalıdır (A–E). "
                "Düz cavabı 3 üsuldan biri ilə qeyd edə bilərsiniz:\n"
                "1) Sual sonunda 'Cavab: B' yazın\n"
                "2) Düz variantın əvvəlinə * qoyun (*B)\n"
                "3) İşarə yoxdursa A standart olaraq düz sayılır\n"
                "Çoxcavablı: 'Cavab: A,C'. Hər sualdan sonra boş sətir buraxın."
            )
            for line in _QUESTION_BANK_TEMPLATE_TXT.splitlines():
                if line.startswith("#") or not line.strip():
                    continue
                doc.add_paragraph(line)

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            response = HttpResponse(
                buf.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = 'attachment; filename="emsarena_test_template.docx"'
            return response
        except Exception:
            logger.exception("Template DOCX generation failed for slug=%s", slug)
            # DOCX alınmadısa TXT-ə düşürük
            file_format = "txt"

    # TXT default
    from django.http import HttpResponse

    response = HttpResponse(_QUESTION_BANK_TEMPLATE_TXT, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = 'attachment; filename="emsarena_test_template.txt"'
    return response


@login_required
@require_POST
def ai_generate_question_bank(request, slug):
    _ensure_teacher(request.user)
    exam = get_teacher_exam_or_404(request, slug=slug)

    if exam.exam_type not in {"test", "written", "coding"}:
        return JsonResponse(
            {
                "ok": False,
                "error": pgettext(
                    "exams.view.question_bank.ai.error",
                    "Bu imtahan tipi üçün AI sual yaratma dəstəklənmir.",
                ),
            },
            status=400,
        )

    source_text = (request.POST.get("source_text") or "").strip()
    uploaded = request.FILES.get("source_file") or request.FILES.get("ai_source_file")
    if uploaded:
        try:
            extracted_text = extract_text_from_upload(uploaded)
        except Exception as exc:
            return JsonResponse(
                {
                    "ok": False,
                    "error": pgettext("exams.view.question_bank.message", "file_read_failed").format(error=exc),
                },
                status=400,
            )
        source_text = "\n\n".join(part for part in [source_text, extracted_text] if part.strip())

    ai_exam_type = "written" if exam.exam_type == "coding" else exam.exam_type
    try:
        result = generate_question_bank_text(
            exam_title=exam.title,
            exam_type=ai_exam_type,
            prompt_text=request.POST.get("prompt", ""),
            source_text=source_text,
            question_count=request.POST.get("question_count") or 5,
            difficulty=request.POST.get("difficulty") or "medium",
            block_name=request.POST.get("block_name") or "",
            language_code=request.LANGUAGE_CODE,
            user_id=request.user.pk,
        )
    except Exception:
        logger.exception("AI question bank endpoint failed for exam %s", exam.pk)
        return JsonResponse(
            {
                "ok": False,
                "error": pgettext(
                    "exams.view.question_bank.ai.error",
                    "AI sual yaratma alınmadı. Bir az sonra yenidən yoxlayın.",
                ),
            },
            status=500,
        )
    status = 200 if result.get("ok") else 400
    return JsonResponse(result, status=status)


@login_required
def create_question_bank(request, slug):
    _ensure_teacher(request.user)
    exam = get_teacher_exam_or_404(request, slug=slug)
    navigation_from_section, navigation_return_to, navigation_query = _resolve_question_bank_navigation(request)

    # Mövcud blokları gətiririk ki, ekranda görsənsin
    blocks = exam.question_blocks.all().order_by("order")

    # Hər blok üçün sualları mətn formatına çeviririk (Textarea üçün)
    # Məsələn: [ {block_obj: block, text_content: "1. Salam\n2. Necəsən"}, ... ]
    blocks_data = []
    for block in blocks:
        questions = block.questions.all().order_by("order")
        # Sualları "1. Sual mətni" formatında birləşdiririk
        text_content = "\n".join([f"{q.order}. {q.text}" for q in questions])

        blocks_data.append(
            {
                "obj": block,
                "text_content": text_content,
                "paint_enabled": block.enable_paint if block.enable_paint is not None else exam.enable_paint,
            }
        )

    return render(
        request,
        "exams/teacher/create_question_bank.html",
        {
            "exam": exam,
            "blocks_data": blocks_data,
            "question_bank_navigation_query": navigation_query,
            "navigation_from_section": navigation_from_section,
            "navigation_return_to": navigation_return_to,
            **_question_bank_title_context(exam),
        },
    )


@login_required
def process_question_bank(request, slug):
    _ensure_teacher(request.user)
    exam = get_teacher_exam_or_404(request, slug=slug)
    _, _, navigation_query = _resolve_question_bank_navigation(request)

    if request.method == "POST":
        # 1. Silinməli olan blokları silirik
        deleted_ids = request.POST.get("deleted_block_ids", "").split(",")
        for d_id in deleted_ids:
            if d_id.strip():
                QuestionBlock.objects.filter(id=d_id, exam=exam).delete()

        # 2. Ümumi sual sayını yenilə
        random_count = _optional_non_negative_int(request.POST.get("random_question_count"))
        if random_count is not None:
            exam.random_question_count = random_count
            exam.save()

        # Adların təkrar olub-olmadığını yoxlamaq üçün set
        used_names = set()

        # ✅ Order hesablamaq üçün counter
        current_order = 1

        # 3. Blokları emal edirik
        for key, value in request.POST.items():
            if key.startswith("block_name_"):
                ui_id = key.split("_")[-1]
                block_name = value.strip()

                # Validation: Eyni sorğuda dublikat ad varmı?
                if block_name.lower() in used_names:
                    messages.error(
                        request,
                        pgettext("exams.view.question_bank.message", "duplicate_block_name_request").format(
                            block_name=block_name
                        ),
                    )
                    return redirect(
                        _append_navigation_query(
                            reverse("exams:create_question_bank", kwargs={"slug": exam.slug}),
                            navigation_query,
                        )
                    )
                used_names.add(block_name.lower())

                content_key = f"block_content_{ui_id}"
                content_text = request.POST.get(content_key, "")
                time_key = f"block_time_{ui_id}"
                time_val = _optional_non_negative_int(request.POST.get(time_key))
                block_paint_enabled = request.POST.get(f"block_enable_paint_{ui_id}") == "on"
                db_id_key = f"block_db_id_{ui_id}"
                db_id = request.POST.get(db_id_key)

                # Validation: Bazada başqa blok eyni adda varmı? (özü xaric)
                existing_check = QuestionBlock.objects.filter(exam=exam, name__iexact=block_name)
                if db_id:
                    existing_check = existing_check.exclude(id=db_id)

                if existing_check.exists():
                    messages.error(
                        request,
                        pgettext("exams.view.question_bank.message", "block_name_exists_db").format(
                            block_name=block_name
                        ),
                    )
                    return redirect(
                        _append_navigation_query(
                            reverse("exams:create_question_bank", kwargs={"slug": exam.slug}),
                            navigation_query,
                        )
                    )

                if block_name:
                    # Blok Yaradılması/Yenilənməsi
                    if db_id:
                        # Bazada yoxlayırıq ki, silinməyibsə (concurrency üçün)
                        block_qs = QuestionBlock.objects.filter(id=db_id, exam=exam)
                        if block_qs.exists():
                            block = block_qs.first()
                            block.name = block_name
                            block.time_limit_minutes = time_val
                            block.enable_paint = block_paint_enabled
                            block.order = current_order  # ✅ Düzgün order
                            block.save()
                        else:
                            continue  # Blok tapılmadısa keçirik
                    else:
                        block = QuestionBlock.objects.create(
                            exam=exam,
                            name=block_name,
                            time_limit_minutes=time_val,
                            enable_paint=block_paint_enabled,
                            order=current_order,  # ✅ Düzgün order (ui_id deyil)
                        )

                    # ✅ Növbəti blok üçün order artır
                    current_order += 1

                    # Sualların Parse edilməsi
                    questions = _parse_written_questions(content_text) if content_text.strip() else []
                    _sync_written_block_questions(block, questions)

        if exam.exam_type == "coding":
            sync_coding_questions_for_exam(exam)

        from apps.exams.services.difficulty import schedule_ai_question_difficulty_warmup

        schedule_ai_question_difficulty_warmup(exam, force=True)

        messages.success(request, pgettext_lazy("exams.view.question_bank.message", "bank_saved"))
        return redirect(
            _append_navigation_query(
                reverse("exams:teacher_exam_detail", kwargs={"slug": exam.slug}),
                navigation_query,
            )
        )

    return redirect(
        _append_navigation_query(
            reverse("exams:create_question_bank", kwargs={"slug": exam.slug}),
            navigation_query,
        )
    )


@login_required
def test_question_bank(request, slug):
    _ensure_teacher(request.user)
    exam = get_teacher_exam_or_404(request, slug=slug)
    navigation_from_section, navigation_return_to, navigation_query = _resolve_question_bank_navigation(request)

    # yalnız test imtahanı üçün
    if exam.exam_type != "test":
        raise Http404()

    blocks = exam.question_blocks.all().order_by("order", "id")

    raw_text = ""
    parsed = []
    selected = set()

    warning_count = 0
    duplicate_count = 0
    error_count = 0
    test_level_warnings = []
    category_counts = {
        "errors": 0,
        "warnings": 0,
        "duplicates": 0,
        "structure": 0,
        "balance": 0,
        "clean": 0,
    }

    # >>> YENİ: UI dəyərləri (Preview klikində sıfırlanmasın deyə)
    # NOTE: 0 = hamısı; None/boş = default 10 göstər
    total_q = exam.questions.filter(is_active=True).count()
    exam_rq = getattr(exam, "random_question_count", None)
    rq_default = min(10, total_q) if exam_rq is None else exam_rq

    exam_dp = getattr(exam, "default_question_points", None) or 1
    dp_default = exam_dp

    # GET-də və POST-da input-ların value-ları buradan gedəcək
    rq_value = str(rq_default)
    dp_value = str(dp_default)

    def build_fp_from_parsed(q):
        return _norm(q["text"]) + "||" + "||".join([_norm(q["options"].get(x, "")) for x in "ABCDE"])

    def build_fp_from_db(eq):
        # DB-də option-lar label saxlamadığı üçün sıra ilə götürürük (A..E)
        opt_map = {}
        opts = list(eq.options.all())
        labels = list("ABCDE")
        for i, opt in enumerate(opts[:5]):
            opt_map[labels[i]] = opt.text
        return _norm(eq.text) + "||" + "||".join([_norm(opt_map.get(x, "")) for x in "ABCDE"])

    def _short_preview(text, length=60):
        """DB dublikat referansı üçün qısa preview (sual nömrəsi yanında göstərmək üçün)."""
        clean = (text or "").strip().replace("\n", " ")
        if len(clean) <= length:
            return clean
        return clean[: length - 1].rstrip() + "…"

    # GET
    if request.method != "POST":
        return render(
            request,
            "exams/teacher/test_question_bank.html",
            {
                "exam": exam,
                "blocks": blocks,
                "raw_text": raw_text,
                "parsed": parsed,
                "selected": selected,
                "warning_count": warning_count,
                "duplicate_count": duplicate_count,
                "error_count": error_count,
                "test_level_warnings": test_level_warnings,
                "category_counts": category_counts,
                # >>> YENİ: input-ların value-ları
                "rq_value": rq_value,
                "dp_value": dp_value,
                "question_bank_navigation_query": navigation_query,
                "navigation_from_section": navigation_from_section,
                "navigation_return_to": navigation_return_to,
            },
        )

    # POST
    action = request.POST.get("action", "preview")

    # >>> YENİ: Preview-də də input dəyərlərini saxla (DB-yə yazmadan!)
    rq_post = (request.POST.get("random_question_count") or "").strip()
    dp_post = (request.POST.get("default_points") or "").strip()

    if rq_post != "":
        rq_value = rq_post  # typed dəyər geri qayıtsın
    if dp_post != "":
        dp_value = dp_post  # typed dəyər geri qayıtsın

    # 1) raw_text-i formdan al (save formunda hidden textarea olmalıdır!)
    raw_text = request.POST.get("raw_text", "")

    # 2) fayl varsa onu oxu (paste varsa fallback kimi qalır)
    uploaded = request.FILES.get("upload_file")
    if uploaded:
        try:
            raw_text = extract_text_from_upload(uploaded)
        except Exception as e:
            # burada fallback: textarea-dakı raw_text qalsın
            messages.error(
                request,
                pgettext("exams.view.question_bank.message", "file_read_failed").format(error=e),
            )

    # 3) preview/save üçün parse et
    if action in ("preview", "save", "download_report"):
        parsed = parse_bulk_mcq(raw_text) or []

        # təhlükəsizlik: warnings açarı hər sualda olsun
        for q in parsed:
            q.setdefault("warnings", [])

        # ---- Duplicate check: import daxilində ----
        # Hər bir fingerprint üçün bütün indeksləri yığırıq, sonra hər sualda
        # qarşılıqlı çarpaz-referans qoyuruq.
        fp_groups = {}
        for idx, q in enumerate(parsed, start=1):
            fp = build_fp_from_parsed(q)
            fp_groups.setdefault(fp, []).append(idx)

        for _fp, indices in fp_groups.items():
            if len(indices) < 2:
                continue
            for pos, idx in enumerate(indices):
                # Sual bu group-da nə dərəcədə yerləşir
                others = [i for i in indices if i != idx]
                # Ən yaxşı referans: ilk variant (ən kiçik nömrə)
                primary_ref = others[0]
                primary_preview = _short_preview(parsed[primary_ref - 1].get("text", ""))
                if pos == 0:
                    # ilk dublikat — sonrakı eyni sualı göstər
                    msg = pgettext("exams.view.question_bank.warning", "duplicate_in_import_first").format(
                        index=idx,
                        next_index=others[0],
                        count=len(others),
                        preview=_short_preview(parsed[others[0] - 1].get("text", "")),
                    )
                else:
                    # sonrakı dublikat — əvvəlki ilk variantı göstər
                    msg = pgettext("exams.view.question_bank.warning", "duplicate_in_import").format(
                        index=idx,
                        previous_index=primary_ref,
                        preview=primary_preview,
                    )

                parsed[idx - 1]["warnings"].append(
                    {
                        "type": "duplicate_in_import",
                        "severity": "error",
                        "msg": msg,
                        "ref": primary_ref if pos > 0 else others[0],
                        "all_refs": others,
                    }
                )

        # ---- Duplicate check: DB-də artıq var? ----
        existing = ExamQuestion.objects.filter(exam=exam).prefetch_related("options").order_by("order", "id")
        existing_by_fp = {}
        for eq in existing:
            existing_by_fp[build_fp_from_db(eq)] = eq

        for idx, q in enumerate(parsed, start=1):
            fp = build_fp_from_parsed(q)
            matched = existing_by_fp.get(fp)
            if matched is not None:
                db_order = matched.order or matched.pk
                db_preview = _short_preview(matched.text)
                q["warnings"].append(
                    {
                        "type": "already_in_exam",
                        "severity": "error",
                        "msg": pgettext("exams.view.question_bank.warning", "already_in_exam").format(
                            index=idx,
                            db_index=db_order,
                            preview=db_preview,
                        ),
                        "ref_db_id": matched.pk,
                        "ref_db_order": db_order,
                    }
                )

        # ---- Test miqyasında "yalnız doğru cavabın uzunluğu" pattern-i ----
        # Tək-tək sualda zərərsiz ola bilər, amma testin böyük hissəsində bu pattern varsa
        # cəlbedici (gözəgörünən) ipucu yaranır — ona görə xəbərdarlıq edirik.
        if parsed:
            long_correct_count = 0
            short_correct_count = 0
            applicable_count = 0
            for q in parsed:
                opts = q.get("options", {}) or {}
                correct = [c for c in q.get("correct", []) if c in opts]
                wrong = [lab for lab in opts if lab not in correct]
                if not correct or not wrong:
                    continue
                applicable_count += 1
                lens_correct = [len((opts.get(c) or "").strip()) for c in correct]
                lens_wrong = [len((opts.get(w) or "").strip()) for w in wrong]
                if not lens_correct or not lens_wrong:
                    continue
                max_c = max(lens_correct)
                avg_w = sum(lens_wrong) / max(1, len(lens_wrong))
                if max_c >= 15 and avg_w > 0 and max_c >= avg_w * 1.8:
                    long_correct_count += 1
                if max(lens_wrong) >= 15 and max_c > 0 and max_c <= avg_w * 0.4:
                    short_correct_count += 1

            # Test həcminin ≥40%-də pattern varsa — informativ xəbərdarlıq
            threshold = max(3, int(applicable_count * 0.4))
            test_level_warnings = []
            if applicable_count >= 5 and long_correct_count >= threshold:
                test_level_warnings.append(
                    {
                        "type": "bulk_correct_too_long",
                        "severity": "warning",
                        "msg": pgettext("exams.view.question_bank.warning", "bulk_correct_too_long").format(
                            ratio=int((long_correct_count / applicable_count) * 100),
                        ),
                    }
                )
            if applicable_count >= 5 and short_correct_count >= threshold:
                test_level_warnings.append(
                    {
                        "type": "bulk_correct_too_short",
                        "severity": "warning",
                        "msg": pgettext("exams.view.question_bank.warning", "bulk_correct_too_short").format(
                            ratio=int((short_correct_count / applicable_count) * 100),
                        ),
                    }
                )
        else:
            test_level_warnings = []

        # ---- Seçilən suallar ----
        selected_from_request = _parse_selected_question_indices(request.POST)
        if selected_from_request is None:
            selected = set(range(1, len(parsed) + 1))
        else:
            selected = selected_from_request

        # ---- Hər sual üçün meta xülasə (UI filter & badge üçün) ----
        # Severity prioriteti: error > warning > info > none
        # severity_rank = {"error": 3, "warning": 2, "info": 1, "none": 0}
        category_counts = {
            "errors": 0,
            "warnings": 0,
            "duplicates": 0,
            "structure": 0,  # variant sayı, boş variant, missing option
            "balance": 0,  # uzunluq balansı
            "clean": 0,  # heç bir warning olmayan sual
        }

        for _idx, q in enumerate(parsed, start=1):
            warnings = q.get("warnings") or []
            counts = {"error": 0, "warning": 0, "info": 0}
            dup_refs = []
            types = set()
            for w in warnings:
                sev = w.get("severity", "warning")
                if sev in counts:
                    counts[sev] += 1
                types.add(w.get("type"))
                if w.get("type") == "duplicate_in_import" and w.get("ref"):
                    dup_refs.append({"kind": "import", "index": w["ref"]})
                if w.get("type") == "already_in_exam":
                    dup_refs.append(
                        {
                            "kind": "db",
                            "index": w.get("ref_db_order"),
                            "db_id": w.get("ref_db_id"),
                        }
                    )

            top_sev = "none"
            for sev in ("error", "warning", "info"):
                if counts[sev]:
                    top_sev = sev
                    break

            has_dup = bool({"duplicate_in_import", "already_in_exam"} & types)
            has_structure = bool(
                {"missing_option", "option_count_recommend_5", "option_count_too_low", "empty_option_text"} & types
            )
            has_balance = bool({"correct_too_long", "correct_too_short"} & types)

            q["meta"] = {
                "top_severity": top_sev,
                "error_count": counts["error"],
                "warning_count": counts["warning"],
                "info_count": counts["info"],
                "total_count": counts["error"] + counts["warning"] + counts["info"],
                "has_duplicate": has_dup,
                "has_structure_issue": has_structure,
                "has_balance_issue": has_balance,
                "dup_refs": dup_refs,
                # JS filtering üçün space-separated tag list
                "flags": " ".join(
                    filter(
                        None,
                        [
                            f"sev-{top_sev}" if top_sev != "none" else "sev-clean",
                            "has-dup" if has_dup else "",
                            "has-structure" if has_structure else "",
                            "has-balance" if has_balance else "",
                            "has-error" if counts["error"] else "",
                            "has-warning" if counts["warning"] else "",
                            "has-info" if counts["info"] else "",
                            "is-clean" if not warnings else "",
                        ],
                    )
                ),
            }

            if counts["error"]:
                category_counts["errors"] += 1
            if counts["warning"]:
                category_counts["warnings"] += 1
            if has_dup:
                category_counts["duplicates"] += 1
            if has_structure:
                category_counts["structure"] += 1
            if has_balance:
                category_counts["balance"] += 1
            if not warnings:
                category_counts["clean"] += 1

        # ---- warning sayları (üst panel üçün) ----
        # severity == "info" sayılmır — yumşaq qeyddir
        warning_count = sum(1 for q in parsed for w in q.get("warnings", []) if w.get("severity", "warning") != "info")
        # test miqyaslı warning-lər də sayılsın
        warning_count += sum(1 for w in test_level_warnings if w.get("severity", "warning") != "info")
        duplicate_count = category_counts["duplicates"]
        error_count = category_counts["errors"]

    # 4) REPORT DOWNLOAD
    if action == "download_report":
        try:
            return _build_question_bank_report_xlsx(
                exam=exam,
                raw_text=raw_text,
                parsed=parsed,
                test_level_warnings=test_level_warnings,
                category_counts=category_counts,
                warning_count=warning_count,
                duplicate_count=duplicate_count,
                error_count=error_count,
            )
        except RuntimeError as exc:
            messages.error(request, str(exc))

    # 5) SAVE
    if action == "save":
        # ---- Exam settings: random_question_count + default_points(+ optional default_question_points) ----
        rq_raw = (request.POST.get("random_question_count") or "").strip()
        dp_raw = (request.POST.get("default_points") or "").strip()

        update_fields = []

        # random_question_count: 0 = hamısı, 10 = 10, 1 = 1 və s.
        if rq_raw.isdigit():
            exam.random_question_count = int(rq_raw)
            update_fields.append("random_question_count")

        # default_points: formdan gəlmirsə, exam.default_question_points varsa onu götür, yoxdursa 1
        if dp_raw.isdigit() and int(dp_raw) > 0:
            default_points = int(dp_raw)
        else:
            default_points = getattr(exam, "default_question_points", None) or 1

        # Exam-də də saxla (əgər field varsa) – köhnə məntiqi pozmur
        if hasattr(exam, "default_question_points"):
            exam.default_question_points = default_points
            update_fields.append("default_question_points")

        created_count = 0
        skipped_count = 0
        points_payload = _parse_points_payload(request.POST)

        with transaction.atomic():
            if update_fields:
                exam.save(update_fields=update_fields)

            # ---- blok seçimi / yeni blok ----
            block_id = request.POST.get("block_id")
            new_block_name = (request.POST.get("new_block_name") or "").strip()
            block_obj = None

            if new_block_name:
                max_order = blocks.aggregate(m=Max("order")).get("m") or 0
                block_obj = QuestionBlock.objects.create(exam=exam, name=new_block_name, order=max_order + 1)
            elif block_id:
                block_obj = QuestionBlock.objects.filter(id=block_id, exam=exam).first()

            # ---- order başlanğıcı ----
            start_order = (ExamQuestion.objects.filter(exam=exam).aggregate(m=Max("order")).get("m") or 0) + 1

            question_rows = []
            option_payloads = []

            for idx, q in enumerate(parsed, start=1):
                if idx not in selected:
                    continue

                # minimum şərt: A-D olsun
                if any(x not in q["options"] for x in ["A", "B", "C", "D"]):
                    skipped_count += 1
                    continue

                # per-question points: compact payload first, legacy input fallback second.
                p_value = points_payload.get(str(idx), request.POST.get(f"points_{idx}"))
                p_raw = str(p_value or "").strip()
                points = int(p_raw) if p_raw.isdigit() and int(p_raw) > 0 else default_points

                question_rows.append(
                    ExamQuestion(
                        exam=exam,
                        block=block_obj,
                        text=q["text"],
                        answer_mode=q["answer_mode"],
                        order=start_order,
                        points=points,
                    )
                )
                option_payloads.append((q["options"], set(q["correct"])))
                start_order += 1

            created_questions = ExamQuestion.objects.bulk_create(question_rows, batch_size=100)
            created_count = len(created_questions)

            option_rows = []
            for eq, (options, correct) in zip(created_questions, option_payloads):
                for lab in "ABCDE":
                    if lab in options:
                        option_rows.append(
                            ExamQuestionOption(
                                question=eq,
                                label=lab,
                                text=options[lab],
                                is_correct=(lab in correct),
                            )
                        )

            if option_rows:
                ExamQuestionOption.objects.bulk_create(option_rows, batch_size=500)

        messages.success(
            request,
            pgettext("exams.view.question_bank.message", "questions_added_summary").format(
                created_count=created_count,
                skipped_count=skipped_count,
            ),
        )
        from apps.exams.services.difficulty import schedule_ai_question_difficulty_warmup

        schedule_ai_question_difficulty_warmup(exam, force=True)
        return redirect(
            _append_navigation_query(
                reverse("exams:test_question_bank", kwargs={"slug": exam.slug}),
                navigation_query,
            )
        )

    # PREVIEW və ya parse sonrası eyni səhifəni göstər
    return render(
        request,
        "exams/teacher/test_question_bank.html",
        {
            "exam": exam,
            "blocks": blocks,
            "raw_text": raw_text,
            "parsed": parsed,
            "selected": selected,
            "warning_count": warning_count,
            "duplicate_count": duplicate_count,
            "error_count": error_count,
            "test_level_warnings": test_level_warnings,
            "category_counts": category_counts,
            # >>> YENİ: Preview refresh olsa da input-lar dolu qalsın
            "rq_value": rq_value,
            "dp_value": dp_value,
            "question_bank_navigation_query": navigation_query,
            "navigation_from_section": navigation_from_section,
            "navigation_return_to": navigation_return_to,
        },
    )
