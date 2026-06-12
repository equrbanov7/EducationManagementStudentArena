"""
Müəllim — Sual Bankı kitabxanası (imtahandan asılı olmayan).

- ``question_bank_list``      — köhnə siyahı URL-i; GET profil bölməsinə,
  POST isə yeni bank yaratma endpoint-inə yönəlir.
- ``question_bank_detail``    — bankın suallarının İDARƏETMƏ görünüşü (ortaq
  ``_question_management`` partial): axtarış, status/dil/sıralama filtri, toplu
  və tək aktiv/deaktiv/sil, modal redaktə.
- ``question_bank_bulk_add``  — yalnız toplu sual yükləmə (ortaq workbench).
- ``bank_question_add/edit``  — tək sual üçün modal əlavə/redaktə (AJAX).
- ``exam_bank_picker``        — imtahan üçün bankdan sual seçib snapshot əlavə et.
"""

import logging

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db import transaction
from django.db.models import Count, Q
from django.db.models.functions import Lower
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils.http import url_has_allowed_host_and_scheme, urlencode
from django.utils.translation import pgettext

from apps.exams.constants import DEFAULT_EXAM_LANGUAGE, EXAM_LANGUAGE_CHOICES
from apps.exams.forms import BankQuestionCreateForm
from apps.exams.models import BankQuestion, BankQuestionOption, QuestionBank, QuestionBlock
from apps.exams.services.access_policy import _ensure_teacher
from apps.exams.services.ai_question_generation import generate_question_bank_text
from apps.exams.services.bank_analysis import analyze_bank_questions
from apps.exams.services.bulk_workbench import (
    analyze_mcq_bulk,
    analyze_written_bulk,
    bank_question_fp_map,
    bank_written_text_map,
    parse_points_payload,
    parse_selected_indices,
)
from apps.exams.services.parsing import extract_text_from_upload
from apps.exams.services.question_bank_attach import (
    _question_fingerprint,
    accessible_banks,
    attach_bank_questions_to_exam,
    bank_questions_queryset,
)
from apps.exams.views.shared.tenant import get_teacher_exam_or_404
from core.tenancy import get_request_organization

_DIFFICULTY_CHOICES = (("easy", "Asan"), ("medium", "Orta"), ("hard", "Çətin"))
_ALLOWED_STATUSES = {"all", "active", "inactive"}
_ALLOWED_SORTS = {"newest", "oldest", "az", "za"}
# Picker modalında lazy-scroll səhifə ölçüsü.
_PICKER_PAGE_SIZE = 40

logger = logging.getLogger(__name__)

# Nümunə şablonlar (toplu sual yükləməsi üçün "Necə yazmalı?" endirməsi)
_BANK_TEMPLATE_TEST_TXT = """\
# EMSArena — Test sual bankı şablonu
# Hər sualın 4 və ya 5 variantı olmalıdır (A–E). Düz cavabı 3 üsuldan biri ilə qeyd edin:
#   1) Sual sonunda "Cavab: B"   2) Düz variantın əvvəlinə * qoyun "*B)"   3) İşarə yoxdursa A.
# Çox cavablı: "Cavab: A,C". Hər sualdan sonra boş sətir buraxın.

1. Şəbəkədə məlumat hansı ölçü vahidi ilə ötürülür?
A) Bit
B) Bayt
C) Volt
D) Hertz
Cavab: A

2. Aşağıdakılardan hansı proqramlaşdırma dilidir?
A) Python
*B) JavaScript
C) Word
D) Excel

3. HTML nədir?
A) Proqramlaşdırma dili
B) İşarələmə dili
C) Verilənlər bazası
D) Əməliyyat sistemi
Cavab: B
"""

_BANK_TEMPLATE_WRITTEN_TXT = """\
# EMSArena — Yazılı sual bankı şablonu
# Hər sual yeni sətirdə nömrə ilə başlasın. Variant lazım deyil — yalnız sual mətni.

1. Verilənlər strukturu nədir? İzah edin.
2. Stack və Queue arasındakı fərqi yazın.
3. Binary axtarış alqoritmini addım-addım təsvir edin.
"""


def _normalize_format(value):
    candidate = (value or "test").strip().lower()
    return candidate if candidate in ("test", "written") else "test"


def _is_modal_request(request):
    return (
        request.GET.get("modal") == "1"
        or request.POST.get("modal") == "1"
        or request.headers.get("X-Requested-With") == "XMLHttpRequest"
    )


def _empty_analysis():
    return {
        "parsed": [],
        "category_counts": {"errors": 0, "warnings": 0, "duplicates": 0, "structure": 0, "balance": 0, "clean": 0},
        "warning_count": 0,
        "duplicate_count": 0,
        "error_count": 0,
        "test_level_warnings": [],
    }


# ---------------------------------------------------------------------------
# Bankın siyahısı + yaratma
# ---------------------------------------------------------------------------
@login_required
def question_bank_list(request):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    profile_bank_url = f"{reverse('accounts:profile')}?section=question-bank"

    if request.method == "POST" and (request.POST.get("action") == "create_bank"):
        name = (request.POST.get("name") or "").strip()
        if not name:
            messages.error(request, pgettext("exams.view.bank.message", "Bank adı boş ola bilməz."))
        else:
            QuestionBank.objects.create(
                name=name,
                subject=(request.POST.get("subject") or "").strip(),
                description=(request.POST.get("description") or "").strip(),
                language=(request.POST.get("language") or DEFAULT_EXAM_LANGUAGE).strip().lower(),
                default_question_type=_normalize_format(request.POST.get("default_question_type")),
                organization=organization,
                created_by=request.user,
                is_shared=bool(request.POST.get("is_shared")),
            )
            messages.success(request, pgettext("exams.view.bank.message", "Sual bankı yaradıldı."))
        next_url = (request.POST.get("next") or "").strip()
        if next_url and url_has_allowed_host_and_scheme(
            next_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()
        ):
            return redirect(next_url)
        return redirect(profile_bank_url)

    return redirect(profile_bank_url)


@login_required
def question_bank_update(request, bank_id):
    """Bankın adını/fənnini/dilini/formatını/paylaşımını redaktə et (yalnız sahib)."""
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)
    if bank.created_by_id != request.user.id:
        messages.error(request, pgettext("exams.view.bank.message", "Yalnız bankın sahibi redaktə edə bilər."))
        return redirect("exams:question_bank_detail", bank_id=bank.id)

    if request.method == "POST":
        name = (request.POST.get("name") or "").strip()
        if not name:
            messages.error(request, pgettext("exams.view.bank.message", "Bank adı boş ola bilməz."))
        else:
            bank.name = name
            bank.subject = (request.POST.get("subject") or "").strip()
            bank.language = (request.POST.get("language") or bank.language).strip().lower()
            bank.default_question_type = _normalize_format(
                request.POST.get("default_question_type") or bank.default_question_type
            )
            bank.is_shared = bool(request.POST.get("is_shared"))
            bank.save(update_fields=["name", "subject", "language", "default_question_type", "is_shared", "updated_at"])
            messages.success(request, pgettext("exams.view.bank.message", "Bank yeniləndi."))
        next_url = (request.POST.get("next") or "").strip()
        if next_url and url_has_allowed_host_and_scheme(
            next_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()
        ):
            return redirect(next_url)
    return redirect("exams:question_bank_detail", bank_id=bank.id)


@login_required
def question_bank_delete(request, bank_id):
    """Bütün bankı sil (yalnız sahib)."""
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)
    if bank.created_by_id != request.user.id:
        messages.error(request, pgettext("exams.view.bank.message", "Yalnız bankın sahibi silə bilər."))
        return redirect("exams:question_bank_detail", bank_id=bank.id)
    if request.method == "POST":
        bank.delete()
        messages.success(request, pgettext("exams.view.bank.message", "Sual bankı silindi."))
        next_url = (request.POST.get("next") or "").strip()
        if next_url and url_has_allowed_host_and_scheme(
            next_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()
        ):
            return redirect(next_url)
        return redirect(f"{reverse('accounts:profile')}?section=question-bank")
    return redirect("exams:question_bank_detail", bank_id=bank.id)


# ---------------------------------------------------------------------------
# Bankın idarəetmə görünüşü (ortaq partial)
# ---------------------------------------------------------------------------
@login_required
def question_bank_detail(request, bank_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)

    # ---- POST: toplu/tək aktiv-deaktiv-sil ----
    if request.method == "POST":
        action = (request.POST.get("bulk_action") or "").strip().lower()

        # Dil üzrə bütün sualları sil (yalnız bank sahibi)
        if action == "delete_language":
            if bank.created_by_id != request.user.id:
                messages.error(
                    request, pgettext("exams.view.bank.message", "Yalnız bankın sahibi bu əməliyyatı edə bilər.")
                )
            else:
                lang = (request.POST.get("language") or "").strip().lower()
                qs = bank.library_questions.all()
                if lang:
                    qs = qs.filter(language=lang)
                deleted = qs.count()
                qs.delete()
                messages.success(
                    request, pgettext("exams.view.bank.message", "{count} sual silindi.").format(count=deleted)
                )
            redirect_params = {}
            keep_lang = (request.POST.get("language") or "").strip()
            if keep_lang:
                redirect_params["language"] = keep_lang
            redirect_url = reverse("exams:question_bank_detail", kwargs={"bank_id": bank.id})
            if redirect_params:
                redirect_url = f"{redirect_url}?{urlencode(redirect_params)}"
            return redirect(redirect_url)

        selected_ids = [int(item) for item in request.POST.getlist("selected_question_ids") if item.isdigit()]
        selected_qs = bank.library_questions.filter(id__in=selected_ids)
        count = selected_qs.count()

        if count == 0:
            messages.warning(request, pgettext("exams.view.bank.message", "Ən azı bir sual seçin."))
        elif action == "deactivate":
            updated = selected_qs.update(is_active=False)
            messages.success(
                request, pgettext("exams.view.bank.message", "{count} sual deaktiv edildi.").format(count=updated)
            )
        elif action == "activate":
            updated = selected_qs.update(is_active=True)
            messages.success(
                request, pgettext("exams.view.bank.message", "{count} sual aktiv edildi.").format(count=updated)
            )
        elif action == "delete":
            selected_qs.delete()
            messages.success(request, pgettext("exams.view.bank.message", "{count} sual silindi.").format(count=count))
        else:
            messages.error(request, pgettext("exams.view.bank.message", "Yanlış əməliyyat."))

        redirect_params = {}
        for key in ("q", "status", "sort", "language", "flag", "page"):
            value = (request.POST.get(key) or "").strip()
            if value:
                redirect_params[key] = value
        url = reverse("exams:question_bank_detail", kwargs={"bank_id": bank.id})
        if redirect_params:
            url = f"{url}?{urlencode(redirect_params)}"
        return redirect(url)

    # ---- GET: filtr + siyahı ----
    allowed_flags = {"has-error", "has-warning", "has-dup", "has-structure", "has-balance", "is-clean"}
    search_query = (request.GET.get("q") or "").strip()[:120]
    status_filter = (request.GET.get("status") or "all").strip().lower()
    sort_filter = (request.GET.get("sort") or "newest").strip().lower()
    language_filter = (request.GET.get("language") or "").strip().lower()
    flag_filter = (request.GET.get("flag") or "").strip().lower()
    if status_filter not in _ALLOWED_STATUSES:
        status_filter = "all"
    if sort_filter not in _ALLOWED_SORTS:
        sort_filter = "newest"
    if flag_filter not in allowed_flags:
        flag_filter = ""

    # Keyfiyyət analizi — seçilmiş dil üzrə (dublikat/xəbərdarlıq/struktur/balans).
    analysis = analyze_bank_questions(bank, language=language_filter or None)
    analysis_enabled = analysis.total_analyzed > 0
    active_flag = flag_filter if (flag_filter and analysis_enabled) else ""

    # Excel keyfiyyət hesabatı (test bankı view-dakı builder təkrar istifadə olunur).
    if analysis_enabled and request.GET.get("export") == "report":
        from types import SimpleNamespace

        from apps.exams.views.teacher.question_bank import _build_question_bank_report_xlsx

        try:
            return _build_question_bank_report_xlsx(
                exam=SimpleNamespace(title=bank.name),
                raw_text="",
                parsed=analysis.parsed_for_report,
                test_level_warnings=[],
                category_counts=analysis.category_counts,
                warning_count=analysis.warning_count,
                duplicate_count=analysis.duplicate_count,
                error_count=analysis.error_count,
            )
        except RuntimeError as exc:
            messages.error(request, str(exc))

    questions = bank.library_questions.prefetch_related("options")
    if search_query:
        matched_ids = (
            bank.library_questions.filter(Q(text__icontains=search_query) | Q(options__text__icontains=search_query))
            .values_list("id", flat=True)
            .distinct()
        )
        questions = questions.filter(id__in=matched_ids)
    if status_filter == "active":
        questions = questions.filter(is_active=True)
    elif status_filter == "inactive":
        questions = questions.filter(is_active=False)
    if language_filter:
        questions = questions.filter(language=language_filter)
    if active_flag:
        questions = questions.filter(id__in=analysis.flagged_ids.get(active_flag, set()))

    if sort_filter == "oldest":
        questions = questions.order_by("created_at", "id")
    elif sort_filter == "az":
        questions = questions.order_by(Lower("text"), "id")
    elif sort_filter == "za":
        questions = questions.order_by(Lower("text").desc(), "id")
    else:
        questions = questions.order_by("-created_at", "-id")

    paginator = Paginator(questions, 12)
    page_obj = paginator.get_page(request.GET.get("page"))

    # Cari səhifədəki suallara analiz nəticəsini bağla (badge/warning üçün)
    if analysis_enabled:
        for question in page_obj.object_list:
            entry = analysis.analysis_by_id.get(question.id)
            question.analysis_meta = entry["meta"] if entry else None
            question.analysis_warnings = entry["warnings"] if entry else []

    # Statistika seçilmiş dilə görə dəyişir
    scoped = bank.library_questions.all()
    if language_filter:
        scoped = scoped.filter(language=language_filter)
    total_questions = scoped.count()
    active_questions = scoped.filter(is_active=True).count()
    inactive_questions = max(total_questions - active_questions, 0)

    base_params = {}
    if search_query:
        base_params["q"] = search_query
    if status_filter != "all":
        base_params["status"] = status_filter
    if sort_filter != "newest":
        base_params["sort"] = sort_filter
    if language_filter:
        base_params["language"] = language_filter
    if active_flag:
        base_params["flag"] = active_flag
    filters_params = {key: value for key, value in base_params.items() if key != "flag"}
    profile_bank_url = f"{reverse('accounts:profile')}?section=question-bank"

    context = {
        "exam": None,
        "bank": bank,
        "page_obj": page_obj,
        "search_query": search_query,
        "status_filter": status_filter,
        "sort_filter": sort_filter,
        "language_filter": language_filter,
        "active_flag": active_flag,
        "total_questions": total_questions,
        "active_questions": active_questions,
        "inactive_questions": inactive_questions,
        "pagination_query": urlencode(base_params),
        "filters_query": urlencode(filters_params),
        "analysis_enabled": analysis_enabled,
        "category_counts": analysis.category_counts,
        "error_count": analysis.error_count,
        "warning_count": analysis.warning_count,
        "duplicate_count": analysis.duplicate_count,
        "clean_count": analysis.clean_count,
        "analyzed_total": analysis.total_analyzed,
        "qm_show_report": True,
        "language_choices": EXAM_LANGUAGE_CHOICES,
        "navigation_from_section": "",
        "navigation_return_to": "",
        # Ortaq idarəetmə partial konteksti
        "qm_context": "bank",
        "qm_title": bank.name,
        "qm_subtitle": pgettext(
            "exams.template.question_bank_detail",
            "Bankın suallarını idarə et: axtar, redaktə et, aktiv/deaktiv et, sil.",
        ),
        "qm_base_url": reverse("exams:question_bank_detail", kwargs={"bank_id": bank.id}),
        "qm_back_url": profile_bank_url,
        "qm_back_label": pgettext("exams.template.question_bank_detail", "Banklar"),
        "qm_bulk_add_url": reverse("exams:question_bank_bulk_add", kwargs={"bank_id": bank.id}),
        "qm_bulk_add_label": pgettext("exams.template.question_bank_detail", "Toplu sual əlavə et"),
        "qm_add_single_url": reverse("exams:bank_question_add", kwargs={"bank_id": bank.id}),
        "qm_show_language_filter": True,
        # Bank redaktə/sil (yalnız sahib)
        "qm_is_owner": bank.created_by_id == request.user.id,
        "qm_default_type_choices": QuestionBank.DEFAULT_QUESTION_TYPE_CHOICES,
        "qm_update_url": reverse("exams:question_bank_update", kwargs={"bank_id": bank.id}),
        "qm_delete_url": reverse("exams:question_bank_delete", kwargs={"bank_id": bank.id}),
        # Word export — bankın suallarını import-uyğun .docx kimi endir.
        "qm_word_export_url": reverse("exams:question_bank_word_export", kwargs={"bank_id": bank.id}),
    }
    return render(request, "exams/teacher/question_bank_detail.html", context)


# ---------------------------------------------------------------------------
# Bankın toplu sual yükləməsi (ortaq workbench — yalnız upload)
# ---------------------------------------------------------------------------
@login_required
def question_bank_bulk_add(request, bank_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)

    raw_text = ""
    parsed = []
    selected = set()
    analysis = _empty_analysis()
    q_format = _normalize_format(
        request.POST.get("q_format") or request.GET.get("format") or bank.default_question_type
    )
    selected_language = (request.POST.get("language") or bank.language or DEFAULT_EXAM_LANGUAGE).strip().lower()

    if request.method == "POST":
        action = (request.POST.get("action") or "preview").strip()
        if action in ("preview", "save"):
            raw_text = request.POST.get("raw_text", "")
            uploaded = request.FILES.get("upload_file")
            if uploaded:
                try:
                    raw_text = extract_text_from_upload(uploaded)
                except Exception as exc:  # noqa: BLE001
                    messages.error(
                        request, pgettext("exams.view.bank.message", "Fayl oxunmadı: {error}").format(error=exc)
                    )

            if q_format == "written":
                analysis = analyze_written_bulk(
                    raw_text, existing_text_map=bank_written_text_map(bank, language=selected_language)
                )
            else:
                analysis = analyze_mcq_bulk(
                    raw_text, existing_fp_map=bank_question_fp_map(bank, language=selected_language)
                )
            parsed = analysis["parsed"]

            selected_from_request = parse_selected_indices(request.POST)
            selected = set(range(1, len(parsed) + 1)) if selected_from_request is None else selected_from_request

            if action == "save":
                created_count = _save_bank_questions(
                    bank=bank,
                    parsed=parsed,
                    selected=selected,
                    language=selected_language,
                    q_format=q_format,
                    points_payload=parse_points_payload(request.POST),
                    created_by=request.user,
                )
                messages.success(
                    request,
                    pgettext("exams.view.bank.message", "{count} sual banka əlavə olundu.").format(count=created_count),
                )
                return redirect("exams:question_bank_detail", bank_id=bank.id)

    context = {
        "exam": None,
        "bank": bank,
        "raw_text": raw_text,
        "parsed": parsed,
        "selected": selected,
        "category_counts": analysis["category_counts"],
        "warning_count": analysis["warning_count"],
        "duplicate_count": analysis["duplicate_count"],
        "error_count": analysis["error_count"],
        "test_level_warnings": analysis["test_level_warnings"],
        "rq_value": "",
        "dp_value": "1",
        # workbench konteksti
        "wb_workbench_key": f"bank-{bank.id}",
        "wb_title": bank.name,
        "wb_subtitle": pgettext(
            "exams.template.question_bank_detail", "Sualları yazın və ya fayl yükləyin, önizləyin, yadda saxlayın."
        ),
        "wb_back_url": reverse("exams:question_bank_detail", kwargs={"bank_id": bank.id}),
        "wb_back_label": pgettext("exams.template.question_bank_detail", "Banka qayıt"),
        "wb_show_settings": False,
        # AI bloku imtahan səhifəsi ilə eyni olsun deyə bankda da göstərilir.
        "wb_ai_url": reverse("exams:ai_generate_bank_questions", kwargs={"bank_id": bank.id}),
        "wb_ai_context": q_format,
        "wb_show_language": True,
        "wb_languages": EXAM_LANGUAGE_CHOICES,
        "wb_selected_language": selected_language,
        # Format toggle YOXDUR — bank yaradılarkən seçilmiş tipə (test/yazılı) görə
        # avtomatik açılır, beləcə imtahan test-bankı ilə eyni görünür.
        "wb_show_format": False,
        "wb_format": q_format,
        "wb_show_report": False,
        "wb_templates": [
            {
                "url": reverse("exams:question_bank_template_download", kwargs={"bank_id": bank.id}) + "?format=docx",
                "label": "DOCX",
                "kind": "docx",
            },
            {
                "url": reverse("exams:question_bank_template_download", kwargs={"bank_id": bank.id}) + "?format=txt",
                "label": "TXT",
                "kind": "txt",
            },
        ],
        "wb_save_label": pgettext("exams.template.question_bank_detail", "Seçilmişləri banka əlavə et"),
    }
    return render(request, "exams/teacher/question_bank_bulk_add.html", context)


def _save_bank_questions(*, bank, parsed, selected, language, q_format, points_payload, created_by):
    rows = []
    option_payloads = []
    for index, question in enumerate(parsed, start=1):
        if index not in selected:
            continue
        text = (question.get("text") or "").strip()
        if not text:
            continue
        raw_points = str(points_payload.get(str(index)) or "").strip()
        points = int(raw_points) if raw_points.isdigit() and int(raw_points) > 0 else 1

        if q_format == "written":
            rows.append(
                BankQuestion(
                    bank=bank,
                    text=text,
                    question_type="written",
                    answer_mode="single",
                    difficulty="medium",
                    language=language,
                    points=points,
                    fingerprint=_question_fingerprint(text),
                    created_by=created_by,
                )
            )
            option_payloads.append(None)
        else:
            options = question.get("options") or {}
            if any(label not in options for label in ("A", "B", "C", "D")):
                continue
            rows.append(
                BankQuestion(
                    bank=bank,
                    text=text,
                    question_type="test",
                    answer_mode=question.get("answer_mode", "single"),
                    difficulty="medium",
                    language=language,
                    points=points,
                    fingerprint=_question_fingerprint(text),
                    created_by=created_by,
                )
            )
            option_payloads.append((options, set(question.get("correct") or [])))

    if not rows:
        return 0
    with transaction.atomic():
        created = BankQuestion.objects.bulk_create(rows, batch_size=100)
        option_rows = []
        for bank_question, payload in zip(created, option_payloads):
            if not payload:
                continue
            options, correct = payload
            for label in "ABCDE":
                if label in options:
                    option_rows.append(
                        BankQuestionOption(
                            question=bank_question, label=label, text=options[label], is_correct=(label in correct)
                        )
                    )
        if option_rows:
            BankQuestionOption.objects.bulk_create(option_rows, batch_size=500)
    return len(created)


# ---------------------------------------------------------------------------
# Nümunə şablon endir (bank toplu-yükləməsi üçün — imtahan səhifəsi ilə eyni)
# ---------------------------------------------------------------------------
@login_required
def question_bank_template_download(request, bank_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)

    q_format = _normalize_format(bank.default_question_type)
    template_text = _BANK_TEMPLATE_WRITTEN_TXT if q_format == "written" else _BANK_TEMPLATE_TEST_TXT
    file_format = (request.GET.get("format") or "txt").lower().strip()

    if file_format == "docx":
        try:
            from io import BytesIO

            from docx import Document

            doc = Document()
            doc.add_heading("EMSArena — Sual bankı şablonu", level=1)
            for line in template_text.splitlines():
                if line.startswith("#") or not line.strip():
                    continue
                doc.add_paragraph(line)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            response = HttpResponse(
                buf.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = 'attachment; filename="sual_sablonu.docx"'
            return response
        except Exception:  # noqa: BLE001 — docx yoxdursa TXT-ə düş
            logger.exception("Bank DOCX template generation failed for bank %s", bank.pk)

    response = HttpResponse(template_text, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = 'attachment; filename="sual_sablonu.txt"'
    return response


@login_required
def question_bank_word_export(request, bank_id):
    """
    Bankdakı sualları Word (.docx) faylı kimi export edir.

    Format import parseri ilə uyğundur (düz cavab `*` prefiksi ilə) —
    müəllim faylı redaktə edib yenidən toplu yükləmə ilə import edə bilər.
    İcazə: bankın əlçatan olduğu istifadəçi (accessible_banks → tenant scoped).
    ?language=az kimi parametrlə dil üzrə filtr mümkündür.
    """
    from urllib.parse import quote

    from apps.exams.services.question_word_export import bank_questions_payload, build_questions_docx

    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)

    language = (request.GET.get("language") or "").strip().lower() or None
    payload = bank_questions_payload(bank, language=language)
    if not payload:
        messages.warning(request, pgettext("exams.view.bank.message", "Export üçün aktiv sual tapılmadı."))
        return redirect("exams:question_bank_detail", bank_id=bank.id)

    subtitle_parts = [f"Sual sayı: {len(payload)}"]
    if bank.subject:
        subtitle_parts.insert(0, f"Fənn: {bank.subject}")
    if language:
        subtitle_parts.append(f"Dil: {language.upper()}")

    buffer = build_questions_docx(
        title=f"Sual bankı — {bank.name}",
        subtitle=" · ".join(subtitle_parts),
        questions=payload,
    )
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    safe_name = quote(f"{bank.name}_suallar.docx")
    response["Content-Disposition"] = f"attachment; filename*=UTF-8''{safe_name}"
    return response


# ---------------------------------------------------------------------------
# AI ilə bankı sualları yarat (imtahandan asılı olmayan) — workbench AI bloku
# imtahan səhifəsi ilə eyni olsun deyə bank toplu-yükləməsində də işləyir.
# ---------------------------------------------------------------------------
@login_required
def ai_generate_bank_questions(request, bank_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)

    # Format: AI kartından (q_format) gəlir, yoxsa bankın default tipi.
    q_format = _normalize_format(request.POST.get("q_format") or bank.default_question_type)

    source_text = (request.POST.get("source_text") or "").strip()
    uploaded = request.FILES.get("source_file") or request.FILES.get("ai_source_file")
    if uploaded:
        try:
            extracted_text = extract_text_from_upload(uploaded)
        except Exception as exc:  # noqa: BLE001
            return JsonResponse(
                {
                    "ok": False,
                    "error": pgettext("exams.view.bank.ai.error", "Fayl oxunmadı: {error}").format(error=exc),
                },
                status=400,
            )
        source_text = "\n\n".join(part for part in [source_text, extracted_text] if part.strip())

    try:
        result = generate_question_bank_text(
            exam_title=bank.name,
            exam_type=q_format,
            prompt_text=request.POST.get("prompt", ""),
            source_text=source_text,
            question_count=request.POST.get("question_count") or 5,
            difficulty=request.POST.get("difficulty") or "medium",
            block_name="",
            language_code=request.LANGUAGE_CODE,
            user_id=request.user.pk,
        )
    except Exception:  # noqa: BLE001
        logger.exception("AI bank question endpoint failed for bank %s", bank.pk)
        return JsonResponse(
            {
                "ok": False,
                "error": pgettext(
                    "exams.view.bank.ai.error", "AI sual yaratma alınmadı. Bir az sonra yenidən yoxlayın."
                ),
            },
            status=500,
        )
    return JsonResponse(result, status=200 if result.get("ok") else 400)


# ---------------------------------------------------------------------------
# Tək sual əlavə/redaktə (modal AJAX)
# ---------------------------------------------------------------------------
def _render_bank_question_form_html(request, *, bank, form, editing=False, question=None):
    return render_to_string(
        "exams/teacher/partials/_question_form.html",
        {"bank": bank, "form": form, "editing": editing, "question": question, "is_modal": True},
        request=request,
    )


@login_required
def bank_question_add(request, bank_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)
    is_modal = _is_modal_request(request)
    q_format = _normalize_format(
        request.POST.get("q_format") or request.GET.get("format") or bank.default_question_type
    )

    if request.method == "POST":
        form = BankQuestionCreateForm(
            request.POST, request.FILES, question_type=q_format, default_language=bank.language
        )
        if form.is_valid():
            question = form.save(commit=False)
            question.bank = bank
            question.question_type = q_format
            question.created_by = request.user
            question.fingerprint = _question_fingerprint(question.text)
            if q_format == "test":
                question.answer_mode = form.cleaned_data.get("answer_mode", "single")
            question.save()
            if q_format == "test":
                form.create_options(question)
            if is_modal:
                return JsonResponse({"success": True, "question_id": question.id})
            return redirect("exams:question_bank_detail", bank_id=bank.id)
        if is_modal:
            return JsonResponse(
                {
                    "success": False,
                    "html": _render_bank_question_form_html(request, bank=bank, form=form, editing=False),
                },
                status=400,
            )
    else:
        form = BankQuestionCreateForm(question_type=q_format, default_language=bank.language)

    return render(
        request,
        "exams/teacher/partials/_question_form.html",
        {"bank": bank, "form": form, "editing": False, "is_modal": True, "q_format": q_format},
    )


@login_required
def bank_question_edit(request, bank_id, question_id):
    _ensure_teacher(request.user)
    organization = get_request_organization(request)
    bank = get_object_or_404(accessible_banks(request.user, organization), id=bank_id)
    question = get_object_or_404(bank.library_questions, id=question_id)
    is_modal = _is_modal_request(request)
    q_format = question.question_type if question.question_type in ("test", "written") else "test"

    if request.method == "POST":
        form = BankQuestionCreateForm(request.POST, request.FILES, instance=question, question_type=q_format)
        if form.is_valid():
            updated = form.save(commit=False)
            updated.bank = bank
            updated.question_type = q_format
            updated.fingerprint = _question_fingerprint(updated.text)
            if q_format == "test":
                updated.answer_mode = form.cleaned_data.get("answer_mode", "single")
            updated.save()
            if q_format == "test":
                form.save_options(updated)
            if is_modal:
                return JsonResponse({"success": True, "question_id": updated.id})
            return redirect("exams:question_bank_detail", bank_id=bank.id)
        if is_modal:
            return JsonResponse(
                {
                    "success": False,
                    "html": _render_bank_question_form_html(
                        request, bank=bank, form=form, editing=True, question=question
                    ),
                },
                status=400,
            )
    else:
        form = BankQuestionCreateForm(instance=question, question_type=q_format)

    return render(
        request,
        "exams/teacher/partials/_question_form.html",
        {"bank": bank, "form": form, "editing": True, "question": question, "is_modal": True, "q_format": q_format},
    )


# ---------------------------------------------------------------------------
# İmtahan üçün bankdan sual seç (picker)
# ---------------------------------------------------------------------------
def _exam_compatible_question_type(exam):
    """İmtahanın tipinə uyğun bank sual tipi (uyğunsuz sual əlavə olunmasın)."""
    return "test" if exam.exam_type == "test" else "written"


def _first_or_default_block(exam):
    """Yazılı imtahan üçün birinci blok (yoxdursa yarat).

    Bankdan əlavə edilən suallar birinci bloka düşür ki, müəllim sonradan onları
    redaktorda istədiyi kimi bloklara böləbilsin.
    """
    block = exam.question_blocks.order_by("order", "id").first()
    if block is None:
        block = QuestionBlock.objects.create(
            exam=exam,
            name=pgettext("exams.view.bank.message", "Bölmə 1"),
            order=1,
        )
    return block


def _bank_language_stats(bank, *, question_type=None):
    """Bank üçün dil üzrə statistika: [(code, label, count)], ümumi, dil sayı."""
    labels = dict(EXAM_LANGUAGE_CHOICES)
    rows = (
        bank_questions_queryset(bank, question_type=question_type)
        .values("language")
        .annotate(count=Count("id"))
        .order_by("-count")
    )
    stats = []
    total = 0
    for row in rows:
        code = row["language"] or DEFAULT_EXAM_LANGUAGE
        count = row["count"]
        total += count
        stats.append({"code": code, "label": labels.get(code, code), "count": count})
    return stats, total, len(stats)


@login_required
def exam_bank_picker(request, slug):
    _ensure_teacher(request.user)
    exam = get_teacher_exam_or_404(request, slug=slug)
    organization = get_request_organization(request)
    is_modal = _is_modal_request(request)
    compatible_type = _exam_compatible_question_type(exam)
    # Yalnız imtahan tipinə uyğun banklar görünsün (test→test bankı, yazılı→yazılı).
    banks = (
        accessible_banks(request.user, organization)
        .filter(default_question_type=compatible_type)
        .order_by("-created_at")
    )

    bank_id = (request.GET.get("bank") or request.POST.get("bank") or "").strip()
    selected_bank = banks.filter(id=bank_id).first() if bank_id else None

    if request.method == "POST" and (request.POST.get("action") == "attach"):
        if selected_bank is None:
            if is_modal:
                return JsonResponse(
                    {"success": False, "error": str(pgettext("exams.view.bank.message", "Əvvəlcə bank seçin."))},
                    status=400,
                )
            messages.error(request, pgettext("exams.view.bank.message", "Əvvəlcə bank seçin."))
        else:
            # "Hamısını seç" rejimi: cari filtrlərə uyğun BÜTÜN sualları əlavə et
            # (yalnız görünən 300-ü deyil), istisna edilənləri çıxmaqla.
            if request.POST.get("select_all") == "1":
                excluded = {int(v) for v in request.POST.getlist("excluded_ids") if v.isdigit()}
                matching_qs = bank_questions_queryset(
                    selected_bank,
                    question_type=compatible_type,
                    language=(request.POST.get("language") or "").strip() or None,
                    difficulty=(request.POST.get("difficulty") or "").strip() or None,
                    search=(request.POST.get("q") or "").strip() or None,
                )
                valid_ids = [qid for qid in matching_qs.values_list("id", flat=True) if qid not in excluded]
            else:
                raw_ids = request.POST.getlist("question_ids")
                requested_ids = [int(value) for value in raw_ids if value.isdigit()]
                # Yalnız imtahanla uyğun tipdə (test/yazılı) sualları əlavə et.
                valid_ids = list(
                    bank_questions_queryset(selected_bank, question_type=compatible_type)
                    .filter(id__in=requested_ids)
                    .values_list("id", flat=True)
                )
            # Yazılı/praktiki imtahanda suallar birinci bloka düşür (sonradan müəllim
            # özü bloklara bölə bilər). Test imtahanında blok məcburi deyil.
            attach_block = None if exam.exam_type == "test" else _first_or_default_block(exam)
            created = attach_bank_questions_to_exam(exam, valid_ids, block=attach_block, created_by=request.user)
            success_msg = pgettext("exams.view.bank.message", "{count} sual imtahana əlavə olundu.").format(
                count=len(created)
            )
            # Yönləndirmə: test → test bankı; yazılı/praktiki → blok redaktoru.
            if exam.exam_type == "test":
                redirect_url = reverse("exams:test_question_bank", kwargs={"slug": exam.slug})
            else:
                redirect_url = reverse("exams:create_question_bank", kwargs={"slug": exam.slug})
                # Müəllimə bloklara bölmə imkanı barədə məlumat ver (redaktorda görünür).
                if created and attach_block is not None:
                    messages.info(
                        request,
                        pgettext(
                            "exams.view.bank.message",
                            "{count} sual “{block}” blokuna əlavə olundu. İstəsəniz redaktorda onları "
                            "ayrı bölmələrə bölə bilərsiniz.",
                        ).format(count=len(created), block=attach_block.name),
                    )
            if is_modal:
                return JsonResponse(
                    {"success": True, "count": len(created), "message": success_msg, "redirect_url": redirect_url}
                )
            messages.success(request, success_msg)
            return redirect(redirect_url)

    try:
        page = max(1, int(request.GET.get("page") or 1))
    except (TypeError, ValueError):
        page = 1

    items_mode = request.GET.get("items") == "1"

    questions = []
    total_count = 0
    bank_language_stats = []
    bank_total_questions = 0
    bank_language_count = 0
    has_more = False
    if selected_bank is not None:
        # Variantlar (akkordeon) üçün prefetch — N+1 sorğunun qarşısını alır.
        question_qs = bank_questions_queryset(
            selected_bank,
            question_type=compatible_type,
            language=(request.GET.get("language") or "").strip() or None,
            difficulty=(request.GET.get("difficulty") or "").strip() or None,
            search=(request.GET.get("q") or "").strip() or None,
        ).prefetch_related("options")
        offset = (page - 1) * _PICKER_PAGE_SIZE
        if items_mode:
            # Scroll əlavəsi: stats/count sorğusu YOX — bir əlavə element çəkib
            # növbəti səhifə olub-olmadığını bilirik (daha sürətli).
            window = list(question_qs[offset : offset + _PICKER_PAGE_SIZE + 1])
            questions = window[:_PICKER_PAGE_SIZE]
            has_more = len(window) > _PICKER_PAGE_SIZE
        else:
            bank_language_stats, bank_total_questions, bank_language_count = _bank_language_stats(
                selected_bank, question_type=compatible_type
            )
            total_count = question_qs.count()
            questions = list(question_qs[offset : offset + _PICKER_PAGE_SIZE])
            has_more = total_count > offset + len(questions)

    context = {
        "exam": exam,
        "banks": banks,
        "selected_bank": selected_bank,
        "questions": questions,
        "total_count": total_count,
        "bank_language_stats": bank_language_stats,
        "bank_total_questions": bank_total_questions,
        "bank_language_count": bank_language_count,
        "compatible_type": compatible_type,
        "language_choices": EXAM_LANGUAGE_CHOICES,
        "difficulty_choices": _DIFFICULTY_CHOICES,
        "search_query": (request.GET.get("q") or "").strip(),
        "language_filter": (request.GET.get("language") or "").strip(),
        "difficulty_filter": (request.GET.get("difficulty") or "").strip(),
        "picker_url": reverse("exams:exam_bank_picker", kwargs={"slug": exam.slug}),
        "page": page,
        "next_page": page + 1,
        "has_more": has_more,
    }
    # AJAX rejimləri: yalnız sual elementləri (scroll əlavəsi) və ya tam content (filtr).
    if request.GET.get("items") == "1":
        return render(request, "exams/teacher/partials/_exam_bank_picker_items.html", context)
    if request.GET.get("content") == "1":
        return render(request, "exams/teacher/partials/_exam_bank_picker_content.html", context)
    if is_modal:
        return render(request, "exams/teacher/partials/_exam_bank_picker_modal.html", context)
    return render(request, "exams/teacher/exam_bank_picker.html", context)
