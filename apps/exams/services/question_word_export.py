"""
Sualların Word (.docx) exportu — sual bankı və imtahan sualları üçün.

Format QƏSDƏN import parserinin qəbul etdiyi formatla eynidir (round-trip):

    1. Sual mətni
    A) variant
    *B) düz variant     <- düz cavab * prefiksi ilə
    C) variant

    2. Növbəti sual ...

Beləliklə müəllim export etdiyi faylı (redaktədən sonra) yenidən toplu
yükləmə ilə geri import edə bilər. Yazılı suallar sadəcə nömrələnmiş
mətn kimi yazılır.

Performans: suallar `prefetch_related("options")` ilə gəlir — option-lar
üçün əlavə sorğu yoxdur; sənəd yaddaşda (BytesIO) qurulur.
"""

from io import BytesIO

OPTION_LABELS = "ABCDEFGH"


def _question_lines(index, text, options):
    """
    Bir sualın sətirlərini qaytarır: nömrəli mətn + (varsa) variantlar.
    options: iterable of (label_or_None, text, is_correct)
    """
    lines = [f"{index}. {(text or '').strip()}"]
    for opt_index, (label, opt_text, is_correct) in enumerate(options):
        letter = (label or "").strip().upper() or OPTION_LABELS[opt_index % len(OPTION_LABELS)]
        prefix = "*" if is_correct else ""
        lines.append(f"{prefix}{letter}) {(opt_text or '').strip()}")
    return lines


def build_questions_docx(*, title, questions, subtitle=""):
    """
    Sualları .docx sənədinə yazır və BytesIO qaytarır.

    questions: iterable of dicts:
        {"text": str, "options": [(label, text, is_correct), ...]}
        (options boş ola bilər — yazılı sual)
    """
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    if subtitle:
        doc.add_paragraph(subtitle)
    doc.add_paragraph(
        "Qeyd: düz cavablar variantın əvvəlindəki * işarəsi ilə qeyd olunub. "
        "Bu faylı redaktə edib yenidən toplu yükləmə ilə import edə bilərsiniz."
    )

    for index, question in enumerate(questions, start=1):
        for line_no, line in enumerate(_question_lines(index, question["text"], question["options"])):
            paragraph = doc.add_paragraph(line)
            if line_no == 0:
                # Sual mətni qalın — oxunaqlıq üçün (parser üçün əhəmiyyətsizdir).
                for run in paragraph.runs:
                    run.bold = True
        # Suallar arası boş sətir — import parseri bölgü üçün bundan istifadə edir.
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def bank_questions_payload(bank, *, language=None, only_active=True):
    """
    BankQuestion → export payload. Options tək prefetch ilə yüklənir.
    """
    questions = bank.library_questions.all()
    if only_active:
        questions = questions.filter(is_active=True)
    if language:
        questions = questions.filter(language=language)
    questions = questions.order_by("created_at", "id").prefetch_related("options")

    payload = []
    for question in questions:
        options = [(opt.label, opt.text, opt.is_correct) for opt in question.options.all()]
        payload.append({"text": question.text, "options": options})
    return payload


def exam_questions_payload(exam, *, language=None):
    """
    ExamQuestion → export payload (imtahanın snapshot sualları).
    """
    questions = exam.questions.all()
    if language:
        questions = questions.filter(language=language)
    questions = questions.order_by("order", "id").prefetch_related("options")

    payload = []
    for question in questions:
        options = [(opt.label, opt.text, opt.is_correct) for opt in question.options.all()]
        payload.append({"text": question.text, "options": options})
    return payload


__all__ = ["bank_questions_payload", "build_questions_docx", "exam_questions_payload"]
